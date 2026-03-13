"""
Script para gerar o Manual Pratico do Fluxo de Trabalho Spec-Driven Development.
Saida: docs/manual-pratico-sdd.docx

Uso: python docs/generate_manual_pratico.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os
from datetime import date

# -----------------------------------------------
# Paleta de cores
# -----------------------------------------------
COR_CR       = "1B4F72"   # azul escuro
COR_PRD      = "2471A3"   # azul medio
COR_ARQ      = "145A32"   # verde escuro
COR_SPEC     = "1E8449"   # verde medio
COR_PLANO    = "B7950B"   # amarelo/laranja
COR_CODIGO   = "566573"   # cinza
COR_VALIDAR  = "922B21"   # vermelho
COR_TITULO   = "1A1A2E"   # azul noite (titulos)
COR_SUBTIT   = "16213E"   # azul escuro (sub)
COR_DESTAQUE = "E8F4F8"   # azul muito claro (fundo destaques)
COR_CINZA_BG = "F2F3F4"   # cinza claro (fundo de info boxes)
COR_SKILL    = "5D3FD3"   # roxo (skill/CLAUDE.md)
BRANCO       = "FFFFFF"
PRETO        = "000000"


# -----------------------------------------------
# Helpers de XML / formatacao
# -----------------------------------------------
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", PRETO))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_run_color(run, hex_color: str):
    run.font.color.rgb = RGBColor.from_string(hex_color)


def set_table_no_spacing(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    cellMar = OxmlElement("w:tblCellMar")
    for side in ["top", "bottom", "left", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "80")
        el.set(qn("w:type"), "dxa")
        cellMar.append(el)
    tblPr.append(cellMar)


def add_page_break(doc):
    doc.add_page_break()


def add_spacer(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_heading(doc, text, level=1, color=None):
    if color is None:
        color = COR_TITULO
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_color(run, color)
    return h


def add_info_box(doc, text, bg_color=None, bold=False):
    if bg_color is None:
        bg_color = COR_CINZA_BG
    table = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return table


def add_code_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1E1E2E")
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    set_run_color(run, "A8FF60")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return table


def add_step_box(doc, step_number, title, description, color=None):
    """Caixa de passo numerado com badge colorido + texto."""
    if color is None:
        color = COR_TITULO
    table = doc.add_table(rows=1, cols=2)
    set_table_no_spacing(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Coluna 0: numero
    c0 = table.cell(0, 0)
    set_cell_bg(c0, color)
    c0.width = Inches(0.5)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(str(step_number))
    r0.font.size = Pt(16)
    r0.bold = True
    set_run_color(r0, BRANCO)
    p0.paragraph_format.space_before = Pt(6)
    p0.paragraph_format.space_after = Pt(6)

    # Coluna 1: titulo + descricao
    c1 = table.cell(0, 1)
    set_cell_bg(c1, COR_CINZA_BG)
    c1.width = Inches(6.0)
    p1 = c1.paragraphs[0]
    r_title = p1.add_run(title)
    r_title.bold = True
    r_title.font.size = Pt(11)
    if description:
        r_desc = p1.add_run(f"\n{description}")
        r_desc.font.size = Pt(10)
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(6)

    return table


def add_checklist_table(doc, items, color=None):
    """Cria tabela de checklist com coluna de checkbox + item."""
    if color is None:
        color = "D5F5E3"
    t = doc.add_table(rows=len(items), cols=2)
    for i, item in enumerate(items):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        set_cell_bg(c0, color)
        c0.width = Inches(0.4)
        r0 = c0.paragraphs[0].add_run("[ ]")
        r0.font.size = Pt(10)
        r0.font.name = "Courier New"
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c1, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r1 = c1.paragraphs[0].add_run(item)
        r1.font.size = Pt(10)
    return t


# -----------------------------------------------
# Estilos globais do documento
# -----------------------------------------------
def configure_document(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i, sz in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {i}"]
        h.font.name = "Calibri"
        h.font.size = Pt(sz)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after  = Pt(6)


# -----------------------------------------------
# CAPA
# -----------------------------------------------
def add_cover(doc):
    add_spacer(doc, 5)

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_TITULO)
    cell.width = Inches(6.5)

    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("Manual Pratico")
    r.font.name = "Calibri"
    r.font.size = Pt(26)
    r.bold = True
    set_run_color(r, BRANCO)
    p_title.paragraph_format.space_before = Pt(16)

    p_sub = cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("Spec-Driven Development")
    r2.font.name = "Calibri"
    r2.font.size = Pt(20)
    r2.bold = True
    set_run_color(r2, "85C1E9")

    p_desc = cell.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_desc.add_run("Guia passo a passo para aplicar SDD no seu projeto")
    r3.font.name = "Calibri"
    r3.font.size = Pt(12)
    r3.italic = True
    set_run_color(r3, "BDC3C7")
    p_desc.paragraph_format.space_after = Pt(16)

    add_spacer(doc, 2)

    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p_ver.add_run(f"Versao 1.0  |  {date.today().strftime('%d/%m/%Y')}")
    r4.font.size = Pt(10)
    set_run_color(r4, "717D7E")

    add_page_break(doc)


# -----------------------------------------------
# CAP 1: O QUE E SDD
# -----------------------------------------------
def add_cap1_o_que_e_sdd(doc):
    add_heading(doc, "1. O Que e Spec-Driven Development (SDD)", 1)

    # 1.1 O Conceito
    add_heading(doc, "1.1 O Conceito", 2, COR_SUBTIT)
    doc.add_paragraph(
        "O Spec-Driven Development (SDD) e uma abordagem de desenvolvimento de software em que "
        "a documentacao e criada antes do codigo. Cada funcionalidade ou alteracao comeca com a "
        "definicao escrita do que deve ser construido -- o \"porque\", o \"o que\" e o \"como\" -- "
        "antes que qualquer linha de codigo seja escrita."
    )
    doc.add_paragraph(
        "Em vez de comecar codando e depois documentar (se sobrar tempo), o SDD inverte a ordem: "
        "primeiro voce escreve o que vai construir, depois constroi exatamente o que foi escrito."
    )

    add_spacer(doc)

    # 1.2 Por Que Usar SDD
    add_heading(doc, "1.2 Por Que Usar SDD", 2, COR_SUBTIT)

    beneficios = [
        ("Rastreabilidade total",
         "Toda decisao tem origem documentada. Voce sempre sabe por que algo foi feito assim."),
        ("Menos retrabalho",
         "Mal-entendidos sao resolvidos no papel, nao no codigo. Refatorar um documento e "
         "infinitamente mais barato do que refatorar codigo em producao."),
        ("Onboarding acelerado",
         "Novos membros da equipe (incluindo agentes de IA) entendem o projeto lendo os documentos, "
         "nao vasculhando o codigo-fonte."),
        ("Escopo controlado",
         "A spec e o contrato. Se nao esta na spec, nao esta no escopo. Isso evita o \"scope creep\" "
         "e feature requests nao documentados."),
        ("Decisoes preservadas",
         "Architecture Decision Records (ADRs) registram decisoes tecnicas e suas alternativas "
         "descartadas, evitando que a equipe repita discussoes ja resolvidas."),
    ]

    for titulo, descricao in beneficios:
        p = doc.add_paragraph(style="List Bullet")
        r_bold = p.add_run(f"{titulo}: ")
        r_bold.bold = True
        r_bold.font.size = Pt(11)
        p.add_run(descricao).font.size = Pt(11)

    add_spacer(doc)

    # 1.3 A Regra de Ouro
    add_heading(doc, "1.3 A Regra de Ouro", 2, COR_SUBTIT)
    add_info_box(
        doc,
        "Regra de Ouro: Documentacao PRIMEIRO -> Codigo DEPOIS.\n"
        "Se voce nao consegue escrever o que vai construir, ainda nao entende o problema bem o suficiente.",
        COR_DESTAQUE,
        bold=True
    )
    add_spacer(doc)

    doc.add_paragraph(
        "Na pratica, isso significa:"
    )
    fluxos = [
        "Novas features: PRD -> Arquitetura -> Spec -> Plano -> Implementacao",
        "Alteracoes/Correcoes: CR -> Avaliar impacto -> Atualizar docs afetados -> Implementar",
        "Bug fix simples: CR -> Implementar -> Atualizar testes",
    ]
    for f in fluxos:
        p = doc.add_paragraph(f, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_spacer(doc)

    # 1.4 Visao Geral do Fluxo
    add_heading(doc, "1.4 Visao Geral do Fluxo", 2, COR_SUBTIT)
    doc.add_paragraph(
        "O fluxo SDD e composto por 8 fases numeradas de 0 a 7. Cada fase tem um documento "
        "correspondente e um proposito especifico."
    )

    add_spacer(doc)

    # Diagrama de fases
    fases = [
        (COR_CR,     "Fase 0\nCR"),
        (COR_PRD,    "Fase 1\nPRD"),
        (COR_ARQ,    "Fase 2\nArquit."),
        (COR_SPEC,   "Fase 3\nSpec"),
        (COR_PLANO,  "Fase 4\nPlano"),
        (COR_CODIGO, "Fase 5\nCodigo"),
        (COR_VALIDAR,"Fase 6\nValid."),
        ("1A5276",   "Fase 7\nDeploy"),
    ]

    table = doc.add_table(rows=2, cols=len(fases))
    set_table_no_spacing(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (cor, label) in enumerate(fases):
        cell = table.cell(0, i)
        set_cell_bg(cell, cor)
        cell.width = Inches(0.85)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.size = Pt(8.5)
        r.bold = True
        set_run_color(r, BRANCO)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    arrow_texts = ["->", "->", "->", "->", "->", "->", "->", ""]
    for i, arrow in enumerate(arrow_texts):
        cell = table.cell(1, i)
        set_cell_bg(cell, "FAFAFA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(arrow)
        r.font.size = Pt(14)
        set_run_color(r, "808B96")

    add_spacer(doc)

    # Tabela de referencia das fases
    headers = ["Fase", "Documento", "Caminho", "Quando Usar"]
    rows = [
        ["0", "Change Request (CR)", "/docs/changes/CR-XXX.md",
         "Alteracoes e correcoes em funcionalidades existentes"],
        ["1", "PRD", "/docs/01-PRD.md",
         "Definicao inicial ou adicao de modulos grandes"],
        ["2", "Arquitetura", "/docs/02-ARCHITECTURE.md",
         "Decisoes de stack, estrutura e padroes"],
        ["3", "Spec Tecnica", "/docs/03-SPEC.md",
         "Detalhamento tecnico de cada feature"],
        ["4", "Plano de Implementacao", "/docs/04-IMPLEMENTATION-PLAN.md",
         "Ordem e dependencias das tarefas"],
        ["5", "Implementacao", "Codigo-fonte",
         "Construcao efetiva"],
        ["6", "Validacao", "Checklist \"Done When\"",
         "Verificar criterios de aceite antes do deploy"],
        ["7", "Deploy e Release", "/docs/05-DEPLOY-GUIDE.md",
         "Procedimentos de deploy, rollback e verificacao"],
    ]

    t = doc.add_table(rows=1 + len(rows), cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        set_run_color(r, BRANCO)

    cores_fase = [COR_CR, COR_PRD, COR_ARQ, COR_SPEC, COR_PLANO, "44546A", COR_VALIDAR, "1A5276"]
    for i, (row, cor) in enumerate(zip(rows, cores_fase)):
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            if j == 0:
                set_cell_bg(cell, cor)
                r = cell.paragraphs[0].add_run(val)
                r.bold = True
                r.font.size = Pt(11)
                set_run_color(r, BRANCO)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
                set_cell_bg(cell, bg)
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(10)

    add_spacer(doc)

    # 1.5 Os Tres Tipos de Trabalho
    add_heading(doc, "1.5 Os Tres Tipos de Trabalho", 2, COR_SUBTIT)
    doc.add_paragraph(
        "Todo trabalho no fluxo SDD se enquadra em uma de tres categorias. "
        "Identificar o tipo correto e o primeiro passo antes de qualquer acao."
    )

    add_spacer(doc)

    # Tabela de decisao
    t2 = doc.add_table(rows=5, cols=3)
    set_table_no_spacing(t2)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER

    labels_tipo = [
        (COR_PRD,     "NOVA FEATURE"),
        (COR_CR,      "ALTERACAO / CORRECAO"),
        (COR_VALIDAR, "BUG SIMPLES"),
    ]
    passos = [
        ["PRD", "Change Request (CR)", "Change Request (CR)"],
        ["Arquitetura", "Avaliar impacto nos docs", "Implementar"],
        ["Spec Tecnica", "Atualizar docs afetados", "Atualizar testes"],
        ["Plano -> Codigo", "Implementar -> Testar", "--"],
    ]

    for j, (cor, label) in enumerate(labels_tipo):
        cell = t2.cell(0, j)
        set_cell_bg(cell, cor)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        set_run_color(r, BRANCO)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    cores_linha = ["EBF5FB", "D6EAF8", "AED6F1", "85C1E9"]
    for i, (passo_row, cor_bg) in enumerate(zip(passos, cores_linha)):
        for j, passo in enumerate(passo_row):
            cell = t2.cell(i + 1, j)
            set_cell_bg(cell, cor_bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(passo)
            r.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_info_box(
        doc,
        "IMPORTANTE: O Change Request (CR) e OBRIGATORIO.\n"
        "Nunca implemente uma feature ou alteracao significativa sem criar o CR primeiro.\n"
        "Mesmo para mudancas urgentes ou aparentemente simples.",
        "FBEEE6",
        bold=True
    )
    add_spacer(doc)


# -----------------------------------------------
# CAP 2: PREPARANDO O AMBIENTE
# -----------------------------------------------
def add_cap2_ambiente(doc):
    add_heading(doc, "2. Preparando o Ambiente", 1)

    # 2.1 Pre-requisitos
    add_heading(doc, "2.1 Pre-requisitos", 2, COR_SUBTIT)
    doc.add_paragraph(
        "Antes de comecar, certifique-se de ter as seguintes ferramentas instaladas:"
    )

    prereqs = [
        ("Claude Code", "Agente de IA que executa o fluxo SDD automaticamente"),
        ("Git", "Controle de versao para gerenciar branches e commits"),
        ("Editor de codigo", "VS Code recomendado, com terminal integrado"),
        ("Python 3.8+", "Necessario apenas se quiser gerar os manuais em .docx"),
    ]

    t = doc.add_table(rows=len(prereqs), cols=2)
    t.style = "Table Grid"
    for i, (ferramenta, descricao) in enumerate(prereqs):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg)
        set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(ferramenta)
        r0.bold = True
        r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(descricao)
        r1.font.size = Pt(10)

    add_spacer(doc)

    doc.add_paragraph("Verifique as instalacoes:")
    add_code_box(doc, "claude --version\ngit --version\npython --version")

    add_spacer(doc)

    # 2.2 Criando a Estrutura de Pastas
    add_heading(doc, "2.2 Criando a Estrutura de Pastas", 2, COR_SUBTIT)

    add_step_box(doc, 1, "Crie a pasta do projeto e inicialize o Git",
                 "Escolha um nome descritivo para o seu projeto.", COR_TITULO)
    add_spacer(doc)
    add_code_box(doc, "mkdir meu-projeto\ncd meu-projeto\ngit init")
    add_spacer(doc)

    add_step_box(doc, 2, "Crie as pastas de documentacao",
                 "O fluxo SDD utiliza tres pastas dentro de docs/.", COR_TITULO)
    add_spacer(doc)
    add_code_box(doc, "mkdir -p docs/templates docs/changes")
    add_spacer(doc)

    doc.add_paragraph(
        "Estrutura resultante:"
    )
    add_code_box(doc,
        "meu-projeto/\n"
        "  docs/\n"
        "    templates/     # Templates dos documentos SDD\n"
        "    changes/       # Change Requests (CR-001, CR-002...)\n"
        "  CLAUDE.md        # (sera criado no Cap. 3)\n"
        "  .gitignore"
    )

    add_spacer(doc)

    # 2.3 Copiando os Templates
    add_heading(doc, "2.3 Copiando os Templates", 2, COR_SUBTIT)

    doc.add_paragraph(
        "Copie os 6 templates abaixo para a pasta docs/templates/ do seu projeto. "
        "Voce pode obte-los do repositorio de referencia ou de qualquer projeto que ja usa SDD."
    )

    add_spacer(doc)

    templates = [
        ("00-template-change-request.md", "Change Request", "Alteracoes em funcionalidades existentes", COR_CR),
        ("01-template-prd.md", "PRD", "Requisitos do produto (novo projeto)", COR_PRD),
        ("02-template-architecture.md", "Arquitetura", "Stack, estrutura, padroes e ADRs", COR_ARQ),
        ("03-template-spec.md", "Spec Tecnica", "Detalhamento tecnico de cada feature", COR_SPEC),
        ("04-template-implementation-plan.md", "Plano de Implementacao", "Tarefas sequenciadas com dependencias", COR_PLANO),
        ("CLAUDE-template.md", "CLAUDE.md", "Instrucoes persistentes para o agente de IA", COR_SKILL),
    ]

    t = doc.add_table(rows=1 + len(templates), cols=3)
    t.style = "Table Grid"
    for j, h in enumerate(["Arquivo", "Documento", "Proposito"]):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        set_run_color(r, BRANCO)

    for i, (arquivo, nome, proposito, cor) in enumerate(templates):
        c0, c1, c2 = t.cell(i + 1, 0), t.cell(i + 1, 1), t.cell(i + 1, 2)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg)
        r0 = c0.paragraphs[0].add_run(arquivo)
        r0.font.name = "Courier New"
        r0.font.size = Pt(9)
        set_cell_bg(c1, cor)
        r1 = c1.paragraphs[0].add_run(nome)
        r1.bold = True
        r1.font.size = Pt(10)
        set_run_color(r1, BRANCO)
        set_cell_bg(c2, bg)
        r2 = c2.paragraphs[0].add_run(proposito)
        r2.font.size = Pt(10)

    add_spacer(doc)

    # 2.4 Instalando a Skill
    add_heading(doc, "2.4 Instalando a Skill /sdd-pipeline (Opcional)", 2, COR_SUBTIT)
    doc.add_paragraph(
        "A skill /sdd-pipeline automatiza todo o fluxo SDD. Para instala-la:"
    )
    add_spacer(doc)

    add_step_box(doc, 1, "Crie a pasta da skill",
                 None, COR_SKILL)
    add_spacer(doc)
    add_code_box(doc, "mkdir -p .claude/skills/sdd-pipeline")
    add_spacer(doc)

    add_step_box(doc, 2, "Copie o arquivo SKILL.md para a pasta criada",
                 "O conteudo da skill e detalhado no Capitulo 4 deste manual.", COR_SKILL)
    add_spacer(doc)

    add_info_box(
        doc,
        "A skill e opcional. Voce pode seguir o fluxo manualmente (Capitulos 3 e 5). "
        "A skill apenas automatiza o processo, garantindo que nenhuma etapa seja pulada.",
        COR_DESTAQUE
    )
    add_spacer(doc)


# -----------------------------------------------
# CAP 3: INICIANDO UM NOVO PROJETO
# -----------------------------------------------
def add_cap3_novo_projeto(doc):
    add_heading(doc, "3. Iniciando um Novo Projeto (Passo a Passo)", 1)

    doc.add_paragraph(
        "Este capitulo guia voce pela criacao de todos os documentos SDD para um projeto novo. "
        "Siga os passos na ordem indicada."
    )

    add_spacer(doc)

    # Passo 1
    add_step_box(doc, 1, "Criar a pasta do projeto e inicializar o Git",
                 "Se voce ja fez isso no Capitulo 2, pule para o Passo 2.", COR_TITULO)
    add_spacer(doc)
    add_code_box(doc, "mkdir meu-projeto && cd meu-projeto && git init")
    add_spacer(doc)

    # Passo 2
    add_step_box(doc, 2, "Criar a estrutura de documentacao",
                 "Pastas para templates, changes e o .gitignore inicial.", COR_TITULO)
    add_spacer(doc)
    add_code_box(doc, "mkdir -p docs/templates docs/changes")
    add_spacer(doc)

    # Passo 3 — CLAUDE.md
    add_step_box(doc, 3, "Criar o CLAUDE.md",
                 "Abra o template CLAUDE-template.md e crie o arquivo CLAUDE.md na raiz do projeto. "
                 "Substitua todos os [placeholders] pelas informacoes do seu projeto.", COR_SKILL)
    add_spacer(doc)

    doc.add_paragraph("Placeholders principais a preencher:")
    placeholders = [
        ("Nome", "[Nome do Produto]", "O nome do seu sistema ou aplicacao"),
        ("Descricao", "[Descricao resumida...]", "O que o sistema faz e para quem"),
        ("Stack", "[Frontend], [Backend], [BD]", "Tecnologias escolhidas"),
        ("Repositorio", "[URL do repositorio Git]", "Link do GitHub/GitLab"),
    ]
    t = doc.add_table(rows=len(placeholders), cols=3)
    t.style = "Table Grid"
    for i, (campo, placeholder, dica) in enumerate(placeholders):
        c0, c1, c2 = t.cell(i, 0), t.cell(i, 1), t.cell(i, 2)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg)
        set_cell_bg(c1, bg)
        set_cell_bg(c2, bg)
        r0 = c0.paragraphs[0].add_run(campo)
        r0.bold = True
        r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(placeholder)
        r1.font.name = "Courier New"
        r1.font.size = Pt(9)
        r2 = c2.paragraphs[0].add_run(dica)
        r2.font.size = Pt(10)

    add_spacer(doc)
    add_info_box(
        doc,
        "O CLAUDE.md e o cerebro do projeto para o agente de IA. Quanto mais completo, "
        "melhores serao as respostas. Mantenha-o sempre atualizado.",
        COR_DESTAQUE
    )
    add_spacer(doc)

    # Passo 4 — PRD
    add_step_box(doc, 4, "Criar o PRD (01-PRD.md)",
                 "Use o template docs/templates/01-template-prd.md como base. "
                 "O PRD define O QUE construir.", COR_PRD)
    add_spacer(doc)

    doc.add_paragraph("Secoes essenciais a preencher:")
    secoes_prd = [
        "Visao Geral do Produto -- problema que resolve e publico-alvo",
        "Objetivos e Metricas de Sucesso -- KPIs mensuraveis",
        "Personas -- perfis dos usuarios com necessidades e frustracoes",
        "Requisitos Funcionais (RF-001, RF-002...) -- cada feature com prioridade",
        "Requisitos Nao-Funcionais (RNF-001...) -- performance, seguranca",
        "User Stories (US-001...) -- Como [persona], quero [acao], para que [beneficio]",
        "Regras de Negocio (RN-001...) -- restricoes que o sistema deve respeitar",
        "Fora de Escopo -- o que NAO sera construido nesta fase",
    ]
    for s in secoes_prd:
        p = doc.add_paragraph(s, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    add_spacer(doc)
    add_info_box(
        doc,
        "Dica: Dedique tempo aos requisitos funcionais (RF-XXX). Cada RF vira uma feature na "
        "Spec e tarefas no Plano. IDs mal definidos aqui propagam confusao em todos os documentos.",
        COR_DESTAQUE
    )
    add_spacer(doc)

    add_code_box(doc, "Salvar como: docs/01-PRD.md")
    add_spacer(doc)

    # Passo 5 — Arquitetura
    add_step_box(doc, 5, "Criar a Arquitetura (02-ARCHITECTURE.md)",
                 "Use o template docs/templates/02-template-architecture.md. "
                 "A Arquitetura define COMO o sistema e construido.", COR_ARQ)
    add_spacer(doc)

    doc.add_paragraph("Secoes essenciais:")
    secoes_arq = [
        "Stack Tecnologica -- tabela com camada / tecnologia / versao / justificativa",
        "Arquitetura Geral -- diagrama Mermaid do fluxo entre componentes",
        "Estrutura de Pastas -- arvore completa de diretorios",
        "Modelagem de Dados -- diagrama ER + detalhamento de entidades",
        "Padroes e Convencoes -- nomenclatura, Git, API, estilo de codigo",
        "ADRs -- decisoes tecnicas com contexto e alternativas descartadas",
        "Deploy e Infraestrutura -- pipeline CI/CD, health check",
    ]
    for s in secoes_arq:
        p = doc.add_paragraph(s, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    add_spacer(doc)
    add_code_box(doc, "Salvar como: docs/02-ARCHITECTURE.md")
    add_spacer(doc)

    # Passo 6 — Spec
    add_step_box(doc, 6, "Criar a Spec Tecnica (03-SPEC.md)",
                 "Use o template docs/templates/03-template-spec.md. "
                 "A Spec detalha EXATAMENTE o que implementar.", COR_SPEC)
    add_spacer(doc)

    doc.add_paragraph("Secoes essenciais:")
    secoes_spec = [
        "Contratos da API -- tabela de todos os endpoints (metodo, path, auth, body, resposta)",
        "Detalhamento por Feature (RF-XXX) -- arquivos, interfaces, logica, endpoints, banco, validacoes",
        "Componentes de UI -- props, estados (loading, empty, error), comportamentos",
        "Fluxos Criticos -- diagramas sequenciais Mermaid",
        "Casos de Borda -- cenarios nao obvios e comportamento esperado",
        "Plano de Testes -- IDs de teste (BT-XXX, FT-XXX) com cenarios e resultados esperados",
    ]
    for s in secoes_spec:
        p = doc.add_paragraph(s, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    add_spacer(doc)
    add_info_box(
        doc,
        "Dica: Inclua uma tabela-resumo de todos os endpoints ANTES de detalhar cada feature. "
        "Isso da uma visao macro para o desenvolvedor antes de mergulhar nos detalhes.",
        COR_DESTAQUE
    )
    add_spacer(doc)
    add_code_box(doc, "Salvar como: docs/03-SPEC.md")
    add_spacer(doc)

    # Passo 7 — Plano
    add_step_box(doc, 7, "Criar o Plano de Implementacao (04-IMPLEMENTATION-PLAN.md)",
                 "Use o template docs/templates/04-template-implementation-plan.md. "
                 "O Plano traduz a Spec em tarefas concretas e sequenciadas.", COR_PLANO)
    add_spacer(doc)

    doc.add_paragraph("Organize as tarefas em 6 grupos:")

    grupos = [
        ("Grupo 1", "Setup e Infraestrutura", "Configuracao inicial, banco, ORM, env"),
        ("Grupo 2", "Modelos e Entidades", "Migrations, models, types/interfaces"),
        ("Grupo 3", "Logica de Negocio", "Services, validacoes, regras"),
        ("Grupo 4", "API / Controllers", "Endpoints, middlewares, rotas"),
        ("Grupo 5", "UI / Frontend", "Componentes, paginas, formularios"),
        ("Grupo 6", "Testes e Refinamento", "Testes de integracao, E2E, code review"),
    ]

    t = doc.add_table(rows=len(grupos), cols=3)
    t.style = "Table Grid"
    for i, (grupo, nome, desc) in enumerate(grupos):
        c0, c1, c2 = t.cell(i, 0), t.cell(i, 1), t.cell(i, 2)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, COR_PLANO)
        set_cell_bg(c1, bg)
        set_cell_bg(c2, bg)
        r0 = c0.paragraphs[0].add_run(grupo)
        r0.bold = True
        r0.font.size = Pt(10)
        set_run_color(r0, BRANCO)
        r1 = c1.paragraphs[0].add_run(nome)
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(desc)
        r2.font.size = Pt(10)

    add_spacer(doc)

    doc.add_paragraph(
        "Cada tarefa deve ter: ID (T-XXX), descricao, arquivos afetados, "
        "referencia ao RF-XXX, dependencias e criterio Done When."
    )
    add_spacer(doc)
    add_code_box(doc, "Salvar como: docs/04-IMPLEMENTATION-PLAN.md")
    add_spacer(doc)

    # Passo 8 — Branch
    add_step_box(doc, 8, "Criar a branch e comecar a codar",
                 "Com todos os documentos prontos, crie a branch de trabalho.", COR_CODIGO)
    add_spacer(doc)
    add_code_box(doc, "git checkout -b feat/mvp")
    add_spacer(doc)

    add_info_box(
        doc,
        "Pronto! Agora voce tem todos os documentos SDD criados. "
        "Implemente tarefa por tarefa conforme o Plano (04-IMPLEMENTATION-PLAN.md), "
        "verificando o Done When de cada uma antes de avancar.",
        "D5F5E3"
    )
    add_spacer(doc)

    # Resumo visual
    add_heading(doc, "Resumo: Documentos Criados", 3, COR_SUBTIT)
    resumo = [
        ("1", "CLAUDE.md", "CLAUDE-template.md", COR_SKILL),
        ("2", "docs/01-PRD.md", "01-template-prd.md", COR_PRD),
        ("3", "docs/02-ARCHITECTURE.md", "02-template-architecture.md", COR_ARQ),
        ("4", "docs/03-SPEC.md", "03-template-spec.md", COR_SPEC),
        ("5", "docs/04-IMPLEMENTATION-PLAN.md", "04-template-implementation-plan.md", COR_PLANO),
    ]

    t = doc.add_table(rows=1 + len(resumo), cols=3)
    t.style = "Table Grid"
    for j, h in enumerate(["#", "Documento Criado", "Template Usado"]):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        set_run_color(r, BRANCO)

    for i, (num, doc_name, template, cor) in enumerate(resumo):
        c0, c1, c2 = t.cell(i + 1, 0), t.cell(i + 1, 1), t.cell(i + 1, 2)
        set_cell_bg(c0, cor)
        r0 = c0.paragraphs[0].add_run(num)
        r0.bold = True
        r0.font.size = Pt(11)
        set_run_color(r0, BRANCO)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c1, bg)
        r1 = c1.paragraphs[0].add_run(doc_name)
        r1.font.name = "Courier New"
        r1.font.size = Pt(9)
        set_cell_bg(c2, bg)
        r2 = c2.paragraphs[0].add_run(template)
        r2.font.name = "Courier New"
        r2.font.size = Pt(9)

    add_spacer(doc)


# -----------------------------------------------
# CAP 4: USANDO A SKILL /SDD-PIPELINE
# -----------------------------------------------
def add_cap4_skill(doc):
    add_heading(doc, "4. Usando a Skill /sdd-pipeline", 1)

    # 4.1 O Que e a Skill
    add_heading(doc, "4.1 O Que e a Skill", 2, COR_SUBTIT)
    doc.add_paragraph(
        "A skill /sdd-pipeline e um pipeline automatizado que executa todas as etapas do "
        "fluxo SDD sem pular nenhuma. Ela le o CLAUDE.md do projeto, determina o modo de trabalho "
        "(Novo Projeto ou Change Request) e guia o agente por cada passo: documentacao, "
        "implementacao, revisao de seguranca, build, validacao, atualizacao de docs e commit."
    )
    add_spacer(doc)

    add_info_box(
        doc,
        "A skill e opcional. Voce pode seguir o fluxo manualmente (como descrito nos "
        "Capitulos 3 e 5). A skill apenas automatiza o processo, garantindo que "
        "nenhuma etapa seja pulada.",
        COR_DESTAQUE
    )
    add_spacer(doc)

    # 4.2 Como Instalar e Configurar
    add_heading(doc, "4.2 Como Instalar e Configurar", 2, COR_SUBTIT)

    add_step_box(doc, 1, "Criar a pasta da skill no projeto",
                 None, COR_SKILL)
    add_spacer(doc)
    add_code_box(doc, "mkdir -p .claude/skills/sdd-pipeline")
    add_spacer(doc)

    add_step_box(doc, 2, "Criar o arquivo SKILL.md",
                 "O arquivo deve conter o frontmatter YAML (name + description) e as instrucoes "
                 "completas do pipeline. Copie de um projeto existente ou do repositorio de referencia.", COR_SKILL)
    add_spacer(doc)

    doc.add_paragraph("Estrutura do SKILL.md:")
    add_code_box(doc,
        "---\n"
        "name: sdd-pipeline\n"
        "description: \"Pipeline completo para implementar features...\"\n"
        "---\n"
        "\n"
        "# /sdd-pipeline -- Pipeline Completo de Feature (SDD)\n"
        "\n"
        "## Instrucoes\n"
        "Execute todas as etapas na ordem indicada...\n"
        "\n"
        "### Passo 0. Ler o CLAUDE.md do projeto\n"
        "### Passo 1. Determinar o Modo de Trabalho\n"
        "## Fluxo A -- Novo Projeto\n"
        "## Fluxo B -- Change Request\n"
        "## Passos Comuns (3 a 8)"
    )
    add_spacer(doc)

    add_step_box(doc, 3, "Invocar a skill no Claude Code",
                 "Basta digitar /sdd-pipeline no inicio de qualquer conversa.", COR_SKILL)
    add_spacer(doc)
    add_code_box(doc, "/sdd-pipeline")
    add_spacer(doc)

    # 4.3 Fluxo A — Novo Projeto
    add_heading(doc, "4.3 Fluxo A -- Novo Projeto", 2, COR_SUBTIT)
    doc.add_paragraph(
        "Use este fluxo quando nao ha base de codigo ou documentacao estabelecida, "
        "ou quando a adicao e grande o suficiente para justificar documentacao de produto."
    )
    add_spacer(doc)

    fluxo_a = [
        ("A1", "Criar o PRD", "Le o template de PRD e cria o documento com todas as secoes. "
         "Valida com o usuario antes de avancar.", COR_PRD),
        ("A2", "Criar a Arquitetura", "Define stack, estrutura de pastas, modelagem de dados, "
         "padroes e ADRs.", COR_ARQ),
        ("A3", "Criar a Spec Tecnica", "Detalha endpoints, tipos, logica de negocio, "
         "validacoes e plano de testes.", COR_SPEC),
        ("A4", "Criar o Plano de Implementacao", "Organiza tarefas em grupos com dependencias "
         "e criterios Done When.", COR_PLANO),
        ("A5", "Criar a Branch", "Cria a branch de trabalho (ex: feat/mvp).", COR_CODIGO),
    ]

    t = doc.add_table(rows=len(fluxo_a), cols=3)
    set_table_no_spacing(t)
    for i, (passo, titulo, desc, cor) in enumerate(fluxo_a):
        c0, c1, c2 = t.cell(i, 0), t.cell(i, 1), t.cell(i, 2)
        set_cell_bg(c0, cor)
        r0 = c0.paragraphs[0].add_run(passo)
        r0.bold = True
        r0.font.size = Pt(11)
        set_run_color(r0, BRANCO)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c1, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r1 = c1.paragraphs[0].add_run(titulo)
        r1.bold = True
        r1.font.size = Pt(10)
        set_cell_bg(c2, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r2 = c2.paragraphs[0].add_run(desc)
        r2.font.size = Pt(10)

    add_spacer(doc)

    # 4.4 Fluxo B — Change Request
    add_heading(doc, "4.4 Fluxo B -- Change Request", 2, COR_SUBTIT)
    doc.add_paragraph(
        "Use este fluxo quando o projeto ja tem base de codigo e documentacao estabelecidos."
    )
    add_spacer(doc)

    fluxo_b = [
        ("B1", "Criar o CR e a Branch",
         "Identifica o proximo numero de CR, le o template, preenche todas as secoes "
         "(resumo, classificacao, AS-IS/TO-BE, impacto, tarefas, criterios de aceite) "
         "e cria a branch feat/CR-XXX-slug.", COR_CR),
        ("B2", "Planejar o Trabalho",
         "Le os documentos existentes (Arquitetura, Spec, Plano), avalia o impacto em CADA "
         "documento e cria um plano detalhado com todas as tarefas.", COR_CR),
    ]

    t = doc.add_table(rows=len(fluxo_b), cols=3)
    set_table_no_spacing(t)
    for i, (passo, titulo, desc, cor) in enumerate(fluxo_b):
        c0, c1, c2 = t.cell(i, 0), t.cell(i, 1), t.cell(i, 2)
        set_cell_bg(c0, cor)
        r0 = c0.paragraphs[0].add_run(passo)
        r0.bold = True
        r0.font.size = Pt(11)
        set_run_color(r0, BRANCO)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c1, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r1 = c1.paragraphs[0].add_run(titulo)
        r1.bold = True
        r1.font.size = Pt(10)
        set_cell_bg(c2, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r2 = c2.paragraphs[0].add_run(desc)
        r2.font.size = Pt(10)

    add_spacer(doc)

    # 4.5 Passos Comuns
    add_heading(doc, "4.5 Passos Comuns (3 a 8)", 2, COR_SUBTIT)
    doc.add_paragraph(
        "Apos concluir o Fluxo A ou B, a skill executa os seguintes passos comuns:"
    )
    add_spacer(doc)

    passos_comuns = [
        ("3", "Implementacao", "Implementa as mudancas conforme a Spec/CR. "
         "Segue convencoes do CLAUDE.md. Escreve testes. Faz commits intermediarios.",
         COR_CODIGO),
        ("4", "Revisao de Seguranca", "Executa checklist OWASP: sem segredos hardcoded, "
         "inputs validados, tokens seguros, ownership verificado, queries parametrizadas, "
         "CORS/headers configurados, dependencias auditadas.",
         COR_VALIDAR),
        ("5", "Verificacao de Build", "Executa comandos de build do projeto (type-check, "
         "lint, testes). Corrige erros antes de prosseguir.",
         COR_PLANO),
        ("6", "Validacao Done When", "Verifica o checklist Done When: funcionalidade ok, "
         "app roda sem erros, testes passando, migration testada.",
         COR_PLANO),
        ("7", "Atualizacao de Documentacao", "Revisa e atualiza TODOS os documentos afetados: "
         "CR, Spec, Plano, Arquitetura, PRD, Deploy Guide, CLAUDE.md.",
         COR_PRD),
        ("8", "Commit, Merge e Push", "Commit final com Conventional Commits, merge na branch "
         "principal com --no-ff, delete da branch e push.",
         "1A5276"),
    ]

    t = doc.add_table(rows=len(passos_comuns), cols=3)
    set_table_no_spacing(t)
    for i, (num, titulo, desc, cor) in enumerate(passos_comuns):
        c0, c1, c2 = t.cell(i, 0), t.cell(i, 1), t.cell(i, 2)
        set_cell_bg(c0, cor)
        r0 = c0.paragraphs[0].add_run(f"Passo {num}")
        r0.bold = True
        r0.font.size = Pt(9)
        set_run_color(r0, BRANCO)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c1, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r1 = c1.paragraphs[0].add_run(titulo)
        r1.bold = True
        r1.font.size = Pt(10)
        set_cell_bg(c2, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r2 = c2.paragraphs[0].add_run(desc)
        r2.font.size = Pt(10)

    add_spacer(doc)

    # 4.6 Dicas de Uso
    add_heading(doc, "4.6 Dicas de Uso", 2, COR_SUBTIT)

    dicas = [
        "Seja especifico no pedido: em vez de \"melhorar a listagem\", escreva \"adicionar "
        "totalizadores por status na tabela de despesas\".",
        "Revise o CR antes de aprovar a implementacao: a skill mostra o CR criado e aguarda "
        "sua confirmacao.",
        "Se a skill perguntar algo, responda antes de continuar: ela pausa em pontos de "
        "ambiguidade para evitar decisoes erradas.",
        "Use /sdd-pipeline no inicio da conversa, antes de qualquer outro pedido.",
    ]
    for d in dicas:
        p = doc.add_paragraph(d, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    add_spacer(doc)


# -----------------------------------------------
# CAP 5: CRIANDO UM CHANGE REQUEST
# -----------------------------------------------
def add_cap5_change_request(doc):
    add_heading(doc, "5. Criando um Change Request (Passo a Passo)", 1)

    # 5.1 Quando Criar um CR
    add_heading(doc, "5.1 Quando Criar um CR", 2, COR_SUBTIT)

    situacoes = [
        "Qualquer alteracao em funcionalidade existente",
        "Correcao de bug nao trivial (que afeta logica ou banco de dados)",
        "Nova feature significativa (antes de criar PRD/Spec)",
        "Refatoracao com impacto em multiplos arquivos",
        "Mudanca de dependencia ou configuracao de ambiente",
    ]
    for s in situacoes:
        p = doc.add_paragraph(s, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_spacer(doc)
    add_info_box(
        doc,
        "NUNCA implemente sem CR. Mesmo para mudancas urgentes ou aparentemente simples. "
        "Se uma alteracao ja foi feita sem CR, crie um retroativamente.",
        "FBEEE6",
        bold=True
    )
    add_spacer(doc)

    # 5.2 Passo 1
    add_step_box(doc, 1, "Identificar o proximo numero de CR",
                 "Liste os CRs existentes para encontrar o proximo numero sequencial.", COR_CR)
    add_spacer(doc)
    add_code_box(doc, "ls docs/changes/\n\n# Exemplo de saida:\n# CR-001-migracao-postgresql.md\n# CR-002-autenticacao-jwt.md\n# CR-003-redesign-frontend.md\n# -> Proximo: CR-004")
    add_spacer(doc)

    # 5.3 Passo 2
    add_step_box(doc, 2, "Criar o arquivo CR a partir do template",
                 "Copie o template e renomeie com o numero e um slug descritivo.", COR_CR)
    add_spacer(doc)
    add_code_box(doc, "cp docs/templates/00-template-change-request.md docs/changes/CR-004-filtro-categoria.md")
    add_spacer(doc)

    # 5.4 Passo 3
    add_step_box(doc, 3, "Preencher cada secao do CR",
                 "Abra o arquivo e preencha todas as secoes. Exemplo abaixo usando um "
                 "cenario ficticio: \"Adicionar filtro por categoria na listagem de despesas\".", COR_CR)
    add_spacer(doc)

    secoes_cr = [
        ("1. Resumo da Mudanca",
         "\"Adicionar filtro por categoria na listagem de despesas mensais, "
         "permitindo que o usuario visualize apenas despesas de uma categoria especifica.\""),
        ("2. Classificacao",
         "Tipo: Nova Feature | Urgencia: Proxima sprint | Complexidade: Media"),
        ("3. Contexto (AS-IS / TO-BE)",
         "AS-IS: A listagem mostra todas as despesas sem filtro.\n"
         "TO-BE: Dropdown de categorias no topo da tabela filtra as despesas."),
        ("4. Detalhamento",
         "Tabela com regras que mudam (ex: endpoint aceita query param ?categoria=X) "
         "e o que NAO muda (ex: ordenacao, paginacao)."),
        ("5. Impacto nos Documentos",
         "Spec: atualizar secao de endpoints. Plano: adicionar tarefas CR-T-01 a CR-T-04."),
        ("6. Impacto no Codigo",
         "Arquivos a modificar: backend/app/routers/expenses.py, "
         "frontend/src/components/ExpenseTable.tsx"),
        ("7. Tarefas de Implementacao",
         "CR-T-01: Adicionar query param no endpoint GET /expenses\n"
         "CR-T-02: Criar componente de filtro no frontend\n"
         "CR-T-03: Testes\n"
         "CR-T-04: Atualizar documentacao"),
        ("8. Criterios de Aceite",
         "[ ] Dropdown exibe todas as categorias\n"
         "[ ] Selecionar categoria filtra a tabela\n"
         "[ ] \"Todas\" remove o filtro"),
        ("9. Riscos",
         "Risco: Performance com muitas categorias. Mitigacao: paginacao no dropdown."),
        ("10. Plano de Rollback",
         "git revert [hash] | Nao ha migration | Nenhuma variavel de ambiente nova"),
    ]

    t = doc.add_table(rows=len(secoes_cr), cols=2)
    t.style = "Table Grid"
    for i, (secao, exemplo) in enumerate(secoes_cr):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg)
        set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(secao)
        r0.bold = True
        r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(exemplo)
        r1.font.size = Pt(9)

    add_spacer(doc)

    # 5.5 Passo 4
    add_step_box(doc, 4, "Criar a branch de trabalho",
                 "Use a convencao feat/ para features e fix/ para correcoes.", COR_CR)
    add_spacer(doc)
    add_code_box(doc, "git checkout -b feat/CR-004-filtro-categoria")
    add_spacer(doc)

    # 5.6 Passo 5
    add_step_box(doc, 5, "Avaliar impacto nos documentos",
                 "Para cada documento, pergunte se precisa de atualizacao.", COR_PRD)
    add_spacer(doc)

    impacto = [
        ("PRD", "Novos requisitos funcionais (RF) ou nao-funcionais (RNF)?"),
        ("Arquitetura", "Nova entidade no banco? Novo padrao ou ADR?"),
        ("Spec Tecnica", "Novos endpoints? Mudanca de contrato existente?"),
        ("Plano de Implementacao", "Novas tarefas a adicionar (secao do CR)?"),
        ("Deploy Guide", "Novas variaveis de ambiente? Nova migration?"),
        ("CLAUDE.md", "Novos arquivos na estrutura? Nova convencao?"),
    ]

    t = doc.add_table(rows=1 + len(impacto), cols=2)
    t.style = "Table Grid"
    for j, h in enumerate(["Documento", "Pergunta a Fazer"]):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        set_run_color(r, BRANCO)

    for i, (doc_name, pergunta) in enumerate(impacto):
        c0, c1 = t.cell(i + 1, 0), t.cell(i + 1, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg)
        set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(doc_name)
        r0.bold = True
        r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(pergunta)
        r1.font.size = Pt(10)

    add_spacer(doc)

    # 5.7 Passo 6
    add_step_box(doc, 6, "Implementar as tarefas do CR",
                 "Implemente uma tarefa por vez, verificando o Done When de cada uma. "
                 "Faca commits intermediarios a cada grupo de tarefas concluido.", COR_SPEC)
    add_spacer(doc)
    add_code_box(doc, "# Apos concluir CR-T-01 e CR-T-02:\ngit add -A\ngit commit -m \"feat: CR-004 - adicionar filtro por categoria no backend e frontend\"")
    add_spacer(doc)

    # 5.8 Passo 7
    add_step_box(doc, 7, "Validar (Done When + Build)",
                 "Verifique todos os criterios de conclusao.", COR_VALIDAR)
    add_spacer(doc)

    validacoes = [
        "Funcionalidade implementada conforme descrito no CR",
        "App roda localmente sem erros (backend + frontend)",
        "Testes existentes continuam passando (sem regressao)",
        "Novos testes cobrem a funcionalidade adicionada",
        "Build passa sem erros de tipo/lint",
    ]
    add_checklist_table(doc, validacoes)
    add_spacer(doc)

    add_code_box(doc, "# Verificar build TypeScript (se aplicavel):\ncd frontend && npx tsc --noEmit -p tsconfig.app.json")
    add_spacer(doc)

    # 5.9 Passo 8
    add_step_box(doc, 8, "Merge e Push",
                 "Faca o merge na branch principal e envie para o repositorio.", "1A5276")
    add_spacer(doc)
    add_code_box(doc,
        "git checkout master\n"
        "git merge feat/CR-004-filtro-categoria --no-ff\n"
        "git branch -d feat/CR-004-filtro-categoria\n"
        "git push origin master"
    )
    add_spacer(doc)


# -----------------------------------------------
# CAP 6: MANTENDO OS DOCUMENTOS SINCRONIZADOS
# -----------------------------------------------
def add_cap6_docs_sync(doc):
    add_heading(doc, "6. Mantendo os Documentos Sincronizados", 1)

    # 6.1 Regra de Ouro
    add_heading(doc, "6.1 A Regra de Ouro da Documentacao", 2, COR_SUBTIT)
    add_info_box(
        doc,
        "A tarefa NAO esta completa ate que os documentos estejam atualizados.\n"
        "Um CLAUDE.md desatualizado faz o agente tomar decisoes baseadas em contexto errado.\n"
        "Uma Spec desatualizada causa implementacoes divergentes do esperado.",
        "FBEEE6",
        bold=True
    )
    add_spacer(doc)

    # 6.2 Quando Atualizar Cada Documento
    add_heading(doc, "6.2 Quando Atualizar Cada Documento", 2, COR_SUBTIT)

    triggers = [
        ("CLAUDE.md", "Ao concluir qualquer CR. Ao criar novos arquivos/pastas. "
         "Ao resolver um problema novo (adicionar ao Troubleshooting)."),
        ("PRD", "Ao adicionar novos requisitos funcionais ou nao-funcionais. "
         "Ao mudar o escopo do produto."),
        ("Arquitetura", "Ao adicionar nova entidade no banco. Ao introduzir nova tecnologia. "
         "Ao tomar decisao tecnica nao obvia (criar ADR)."),
        ("Spec Tecnica", "Ao criar/alterar endpoints. Ao mudar contratos de API. "
         "Ao alterar logica de negocio."),
        ("Plano de Implementacao", "Ao incorporar um novo CR (adicionar secao com tarefas CR-T-XX). "
         "Ao concluir grupos de tarefas (atualizar status)."),
        ("Deploy Guide", "Ao adicionar variaveis de ambiente. Ao criar nova migration. "
         "Ao mudar procedimentos de deploy."),
        ("CR", "Ao concluir implementacao (status -> Concluido). "
         "Ao descobrir riscos nao previstos."),
    ]

    t = doc.add_table(rows=1 + len(triggers), cols=2)
    t.style = "Table Grid"
    for j, h in enumerate(["Documento", "Atualizar Quando..."]):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        set_run_color(r, BRANCO)

    cores_doc = [COR_SKILL, COR_PRD, COR_ARQ, COR_SPEC, COR_PLANO, "1A5276", COR_CR]
    for i, ((doc_name, trigger), cor) in enumerate(zip(triggers, cores_doc)):
        c0, c1 = t.cell(i + 1, 0), t.cell(i + 1, 1)
        set_cell_bg(c0, cor)
        r0 = c0.paragraphs[0].add_run(doc_name)
        r0.bold = True
        r0.font.size = Pt(10)
        set_run_color(r0, BRANCO)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c1, bg)
        r1 = c1.paragraphs[0].add_run(trigger)
        r1.font.size = Pt(10)

    add_spacer(doc)

    # 6.3 Checklist de Revisao (Passo 7 da Skill)
    add_heading(doc, "6.3 Checklist de Revisao de Documentos", 2, COR_SUBTIT)
    doc.add_paragraph(
        "Ao concluir uma implementacao (Passo 7 da skill), revise CADA documento abaixo. "
        "\"Revisar\" significa abrir o documento, verificar se o conteudo reflete as mudancas, "
        "e atualizar o que estiver desatualizado."
    )
    add_spacer(doc)

    checklist_docs = [
        "CR: Status atualizado para \"Concluido\" + entrada no changelog",
        "Spec Tecnica: Descricoes de endpoints/modelos/componentes refletem as mudancas?",
        "Plano de Implementacao: CR referenciado no header e na tabela de visao geral?",
        "Arquitetura: Novas decisoes arquiteturais, padroes ou ADRs necessarios?",
        "PRD: Novos requisitos funcionais ou nao-funcionais?",
        "Deploy Guide: Novas variaveis, migrations ou procedimentos de deploy?",
        "CLAUDE.md: Novos CRs na lista, mudancas estruturais, convencoes alteradas?",
    ]
    add_checklist_table(doc, checklist_docs, "D6EAF8")

    add_spacer(doc)
    add_info_box(
        doc,
        "Se um documento nao precisa de atualizacao, registre a justificativa explicitamente.\n"
        "Exemplo: \"PRD -- nao precisa: CR nao altera requisitos funcionais\".",
        COR_DESTAQUE
    )
    add_spacer(doc)

    # 6.4 Referencia Cruzada de IDs
    add_heading(doc, "6.4 Referencia Cruzada de IDs", 2, COR_SUBTIT)
    doc.add_paragraph(
        "O fluxo SDD usa um sistema de IDs para criar rastreabilidade entre documentos. "
        "Um ID em um documento deve aparecer exatamente igual em todos os outros que "
        "fazem referencia a ele."
    )
    add_spacer(doc)

    ids = [
        ("RF-XXX",  COR_PRD,     "Requisito Funcional",           "PRD -> Spec -> Plano"),
        ("RN-XXX",  COR_ARQ,     "Regra de Negocio",              "PRD -> Spec"),
        ("US-XXX",  COR_SPEC,    "User Story",                    "PRD"),
        ("T-XXX",   COR_PLANO,   "Tarefa de Implementacao",       "Plano de Implementacao"),
        ("CR-XXX",  COR_CR,      "Change Request",                "Todas as fases"),
        ("CR-T-XX", COR_CR,      "Tarefa de um CR",               "Plano (secao do CR)"),
        ("ADR-XXX", COR_ARQ,     "Architecture Decision Record",  "Arquitetura"),
        ("BT-XXX",  COR_VALIDAR, "Backend Test",                  "Spec -> Testes"),
        ("FT-XXX",  COR_VALIDAR, "Full Flow Test",                "Spec -> Testes"),
    ]

    t = doc.add_table(rows=1 + len(ids), cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(["Prefixo", "Tipo", "Usado em", "Fluxo"]):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        set_run_color(r, BRANCO)

    for i, (prefixo, cor, tipo, usado_em) in enumerate(ids):
        c0 = t.cell(i + 1, 0)
        set_cell_bg(c0, cor)
        r0 = c0.paragraphs[0].add_run(prefixo)
        r0.bold = True
        r0.font.size = Pt(10)
        set_run_color(r0, BRANCO)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for j, val in enumerate([tipo, usado_em]):
            cell = t.cell(i + 1, j + 1)
            bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
            set_cell_bg(cell, bg)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10)

        # Coluna "Fluxo" fica vazia na tabela — usado_em ja cobre
        c3 = t.cell(i + 1, 3)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c3, bg)

    add_spacer(doc)
    add_info_box(
        doc,
        "Rastreabilidade: Se o RF-013 aparecer no PRD, ele deve aparecer como referencia "
        "na Spec (secao da feature) e no Plano (coluna Ref das tarefas). "
        "Isso permite rastrear qualquer requisito de ponta a ponta.",
        COR_DESTAQUE
    )
    add_spacer(doc)


# -----------------------------------------------
# CAP 7: CHECKLIST RAPIDO POR CENARIO
# -----------------------------------------------
def add_cap7_checklists(doc):
    add_heading(doc, "7. Checklist Rapido por Cenario", 1)

    doc.add_paragraph(
        "Use estas checklists como referencia rapida para cada tipo de trabalho."
    )
    add_spacer(doc)

    # 7.1 Novo Projeto
    add_heading(doc, "7.1 Quero Criar um Projeto Novo", 2, COR_PRD)
    items_novo = [
        "Criar pasta do projeto + git init",
        "Criar docs/templates/ e docs/changes/",
        "Copiar os 6 templates para docs/templates/",
        "Criar CLAUDE.md na raiz (do template CLAUDE-template.md)",
        "Criar docs/01-PRD.md (do template 01-template-prd.md)",
        "Criar docs/02-ARCHITECTURE.md (do template 02-template-architecture.md)",
        "Criar docs/03-SPEC.md (do template 03-template-spec.md)",
        "Criar docs/04-IMPLEMENTATION-PLAN.md (do template 04-template-implementation-plan.md)",
        "Criar branch feat/mvp e comecar a implementar tarefa por tarefa",
    ]
    add_checklist_table(doc, items_novo, "D5F5E3")
    add_spacer(doc)

    # 7.2 Adicionar Feature
    add_heading(doc, "7.2 Quero Adicionar uma Feature", 2, COR_CR)
    items_feature = [
        "Identificar proximo numero de CR (ls docs/changes/)",
        "Criar CR usando template 00-template-change-request.md",
        "Preencher TODAS as secoes do CR (resumo, AS-IS/TO-BE, impacto, tarefas, aceite)",
        "Criar branch: git checkout -b feat/CR-XXX-slug",
        "Avaliar impacto em cada documento (PRD, Arq, Spec, Plano, Deploy, CLAUDE.md)",
        "Atualizar documentos afetados ANTES de implementar",
        "Implementar tarefas do CR uma por vez (commits intermediarios)",
        "Executar revisao de seguranca (se aplicavel)",
        "Verificar build (tsc --noEmit, testes passando)",
        "Validar criterios de aceite + Done When Universal",
        "Atualizar TODOS os documentos (Passo 7 da skill)",
        "Merge em master (--no-ff), deletar branch, push",
    ]
    add_checklist_table(doc, items_feature, "D6EAF8")
    add_spacer(doc)

    # 7.3 Corrigir Bug
    add_heading(doc, "7.3 Quero Corrigir um Bug", 2, COR_VALIDAR)
    items_bug = [
        "Criar CR (mesmo para bugs simples -- documentar o que estava errado)",
        "Criar branch: git checkout -b fix/CR-XXX-slug",
        "Implementar a correcao",
        "Escrever/atualizar testes que cobrem o cenario do bug",
        "Verificar build e testes passando",
        "Atualizar documentos afetados (Spec, CLAUDE.md)",
        "Merge em master (--no-ff), deletar branch, push",
    ]
    add_checklist_table(doc, items_bug, "FADBD8")
    add_spacer(doc)

    # 7.4 Deploy
    add_heading(doc, "7.4 Quero Fazer Deploy", 2, "1A5276")
    items_deploy = [
        "Verificar que o build passa sem erros (type-check + lint)",
        "Todos os testes passando (sem regressao)",
        "Documentos atualizados e sincronizados",
        "Branch do CR merged em master com --no-ff",
        "git push origin master (dispara auto-deploy)",
        "Verificar health check apos deploy (GET /api/health)",
    ]
    add_checklist_table(doc, items_deploy, "D5F5E3")
    add_spacer(doc)

    # Conventional Commits
    add_heading(doc, "7.5 Referencia: Conventional Commits", 2, COR_SUBTIT)

    commits = [
        ("feat:", "Nova feature implementada", "feat: CR-005 - daily expenses module"),
        ("fix:", "Correcao de bug", "fix: CR-007 - prevent duplicate installments"),
        ("docs:", "Atualizacao de documentacao", "docs: update spec for CR-005"),
        ("refactor:", "Refatoracao sem mudanca de comportamento", "refactor: extract status logic"),
        ("test:", "Adicao ou correcao de testes", "test: add BT-011 to BT-015"),
        ("chore:", "Tarefas de manutencao", "chore: update dependencies"),
    ]

    t = doc.add_table(rows=len(commits), cols=3)
    t.style = "Table Grid"
    for i, (prefixo, desc, exemplo) in enumerate(commits):
        c0, c1, c2 = t.cell(i, 0), t.cell(i, 1), t.cell(i, 2)
        set_cell_bg(c0, COR_TITULO)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c1, bg)
        set_cell_bg(c2, bg)
        r0 = c0.paragraphs[0].add_run(prefixo)
        r0.font.name = "Courier New"
        r0.font.size = Pt(10)
        set_run_color(r0, "85C1E9")
        r1 = c1.paragraphs[0].add_run(desc)
        r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(exemplo)
        r2.font.name = "Courier New"
        r2.font.size = Pt(9)

    add_spacer(doc, 2)

    add_info_box(
        doc,
        "Este manual esta disponivel em: docs/manual-pratico-sdd.docx\n"
        "Os templates estao em: docs/templates/\n"
        "A skill /sdd-pipeline esta em: .claude/skills/sdd-pipeline/SKILL.md\n"
        "O manual conceitual (teoria do SDD) esta em: docs/manual-workflow-sdd.docx",
        COR_CINZA_BG
    )
    add_spacer(doc)


# -----------------------------------------------
# MAIN
# -----------------------------------------------
def main():
    doc = Document()
    configure_document(doc)

    add_cover(doc)
    add_cap1_o_que_e_sdd(doc)
    add_page_break(doc)
    add_cap2_ambiente(doc)
    add_page_break(doc)
    add_cap3_novo_projeto(doc)
    add_page_break(doc)
    add_cap4_skill(doc)
    add_page_break(doc)
    add_cap5_change_request(doc)
    add_page_break(doc)
    add_cap6_docs_sync(doc)
    add_page_break(doc)
    add_cap7_checklists(doc)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, "manual-pratico-sdd.docx")
    doc.save(output_path)
    print(f"[OK]  Manual pratico gerado: {output_path}")


if __name__ == "__main__":
    main()
