"""
Script para gerar o Manual do Fluxo de Trabalho Spec-Driven Development.
Saída: docs/manual-workflow-sdd.docx

Uso: python docs/generate_manual.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os
from datetime import date

# ─────────────────────────────────────────────
# Paleta de cores
# ─────────────────────────────────────────────
COR_CR       = "1B4F72"   # azul escuro
COR_PRD      = "2471A3"   # azul médio
COR_ARQ      = "145A32"   # verde escuro
COR_SPEC     = "1E8449"   # verde médio
COR_PLANO    = "B7950B"   # amarelo/laranja
COR_CODIGO   = "566573"   # cinza
COR_VALIDAR  = "922B21"   # vermelho
COR_TITULO   = "1A1A2E"   # azul noite (títulos)
COR_SUBTIT   = "16213E"   # azul escuro (sub)
COR_DESTAQUE = "E8F4F8"   # azul muito claro (fundo destaques)
COR_CINZA_BG = "F2F3F4"   # cinza claro (fundo de info boxes)
BRANCO       = "FFFFFF"
PRETO        = "000000"


# ─────────────────────────────────────────────
# Helpers de XML / formatação
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """Define cor de fundo de uma célula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Define bordas de uma célula individualmente."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", PRETO))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_run_color(run, hex_color: str):
    run.font.color.rgb = RGBColor.from_string(hex_color)


def set_table_no_spacing(table):
    """Remove espaço interno de células da tabela."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    cellMar = OxmlElement("w:tblCellMar")
    for side in ["top", "bottom", "left", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "80")
        el.set(qn("w:type"), "dxa")
        cellMar.append(el)
    tblPr.append(cellMar)


def add_page_break(doc):
    doc.add_page_break()


def add_spacer(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_heading(doc, text, level=1, color=COR_TITULO):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_color(run, color)
    return h


def add_info_box(doc, text, bg_color=COR_CINZA_BG, bold=False):
    """Caixa de destaque com fundo colorido."""
    table = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return table


def add_code_box(doc, text):
    """Caixa de código (fundo escuro, fonte monoespaçada)."""
    table = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1E1E2E")
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    set_run_color(run, "A8FF60")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return table


# ─────────────────────────────────────────────
# Estilos globais do documento
# ─────────────────────────────────────────────
def configure_document(doc):
    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Estilo Normal
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Estilos de Heading
    for i, sz in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {i}"]
        h.font.name = "Calibri"
        h.font.size = Pt(sz)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after  = Pt(6)


# ─────────────────────────────────────────────
# 1. CAPA
# ─────────────────────────────────────────────
def add_cover(doc):
    add_spacer(doc, 5)

    # Bloco colorido de título
    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_TITULO)
    cell.width = Inches(6.5)

    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("Guia do Fluxo de Trabalho")
    r.font.name = "Calibri"
    r.font.size = Pt(26)
    r.bold = True
    set_run_color(r, BRANCO)
    p_title.paragraph_format.space_before = Pt(16)

    p_sub = cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("Spec-Driven Development")
    r2.font.name = "Calibri"
    r2.font.size = Pt(20)
    r2.bold = True
    set_run_color(r2, "85C1E9")

    p_desc = cell.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_desc.add_run("Um método para desenvolver software com documentação como base")
    r3.font.name = "Calibri"
    r3.font.size = Pt(12)
    r3.italic = True
    set_run_color(r3, "BDC3C7")
    p_desc.paragraph_format.space_after = Pt(16)

    add_spacer(doc, 2)

    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p_ver.add_run(f"Versão 1.0  ·  {date.today().strftime('%d/%m/%Y')}")
    r4.font.size = Pt(10)
    set_run_color(r4, "717D7E")

    add_page_break(doc)


# ─────────────────────────────────────────────
# 2. INTRODUÇÃO
# ─────────────────────────────────────────────
def add_introducao(doc):
    add_heading(doc, "1. Introdução", 1)

    doc.add_paragraph(
        "O Spec-Driven Development (SDD) é uma abordagem de desenvolvimento de software em que "
        "a documentação é criada antes do código. Cada funcionalidade ou alteração começa com a "
        "definição escrita do que deve ser construído — o \"porquê\", o \"o quê\" e o \"como\" — "
        "antes que qualquer linha de código seja escrita."
    )

    add_heading(doc, "Por que documentar antes de codar?", 3, COR_SUBTIT)

    items = [
        ("Rastreabilidade total",
         "Toda decisão tem origem documentada. Você sempre sabe por que algo foi feito assim."),
        ("Menos retrabalho",
         "Mal-entendidos são resolvidos no papel, não no código. Refatorar um documento é infinitamente "
         "mais barato do que refatorar código em produção."),
        ("Onboarding acelerado",
         "Novos membros da equipe (incluindo agentes de IA) entendem o projeto lendo os documentos, "
         "não vasculhando o código-fonte."),
        ("Escopo controlado",
         "A spec é o contrato. Se não está na spec, não está no escopo. Isso evita o \"scope creep\" "
         "e feature requests não documentados."),
        ("Decisões preservadas",
         "Architecture Decision Records (ADRs) registram decisões técnicas e suas alternativas "
         "descartadas, evitando que a equipe repita discussões já resolvidas."),
    ]

    for titulo, descricao in items:
        p = doc.add_paragraph(style="List Bullet")
        r_bold = p.add_run(f"{titulo}: ")
        r_bold.bold = True
        r_bold.font.size = Pt(11)
        p.add_run(descricao).font.size = Pt(11)

    add_spacer(doc)
    add_info_box(
        doc,
        "💡  Regra de Ouro: Documentação PRIMEIRO → Código DEPOIS.\n"
        "Se você não consegue escrever o que vai construir, ainda não entende o problema bem o suficiente.",
        COR_DESTAQUE,
        bold=False
    )
    add_spacer(doc)


# ─────────────────────────────────────────────
# 3. VISÃO GERAL DO FLUXO
# ─────────────────────────────────────────────
def add_visao_geral(doc):
    add_heading(doc, "2. Visão Geral do Fluxo", 1)

    doc.add_paragraph(
        "O fluxo SDD é composto por 8 fases numeradas de 0 a 7. Cada fase tem um documento "
        "correspondente e um propósito específico. O fluxo completo vai da definição do problema "
        "até o deploy em produção."
    )

    add_spacer(doc)

    # Diagrama de fases — tabela colorida 8 células
    fases = [
        (COR_CR,     "Fase 0\nCR"),
        (COR_PRD,    "Fase 1\nPRD"),
        (COR_ARQ,    "Fase 2\nArquit."),
        (COR_SPEC,   "Fase 3\nSpec"),
        (COR_PLANO,  "Fase 4\nPlano"),
        (COR_CODIGO, "Fase 5\nCódigo"),
        (COR_VALIDAR,"Fase 6\nValid."),
        ("1A5276",   "Fase 7\nDeploy"),
    ]

    table = doc.add_table(rows=2, cols=len(fases))
    set_table_no_spacing(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Linha 1: blocos coloridos com número da fase
    for i, (cor, label) in enumerate(fases):
        cell = table.cell(0, i)
        set_cell_bg(cell, cor)
        cell.width = Inches(0.85)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.size = Pt(8.5)
        r.bold = True
        set_run_color(r, BRANCO)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    # Linha 2: setas entre fases
    arrow_texts = ["→", "→", "→", "→", "→", "→", "→", ""]
    for i, arrow in enumerate(arrow_texts):
        cell = table.cell(1, i)
        set_cell_bg(cell, "FAFAFA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(arrow)
        r.font.size = Pt(14)
        set_run_color(r, "808B96")

    add_spacer(doc)

    # Tabela de referência das fases
    add_heading(doc, "Referência das Fases", 3, COR_SUBTIT)

    headers = ["Fase", "Documento", "Caminho", "Quando Usar"]
    rows = [
        ["0", "Change Request (CR)", "/docs/changes/CR-XXX.md",
         "Alterações e correções em funcionalidades existentes"],
        ["1", "PRD", "/docs/01-PRD.md",
         "Definição inicial ou adição de módulos grandes"],
        ["2", "Arquitetura", "/docs/02-ARCHITECTURE.md",
         "Decisões de stack, estrutura e padrões"],
        ["3", "Spec Técnica", "/docs/03-SPEC.md",
         "Detalhamento técnico de cada feature"],
        ["4", "Plano de Implementação", "/docs/04-IMPLEMENTATION-PLAN.md",
         "Ordem e dependências das tarefas"],
        ["5", "Implementação", "Código-fonte",
         "Construção efetiva"],
        ["6", "Validação", "Checklist \"Done When\"",
         "Verificar critérios de aceite antes do deploy"],
        ["7", "Deploy e Release", "/docs/05-DEPLOY-GUIDE.md",
         "Procedimentos de deploy, rollback e verificação"],
    ]

    t = doc.add_table(rows=1 + len(rows), cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabeçalho
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        set_run_color(r, BRANCO)

    # Dados
    cores_fase = [COR_CR, COR_PRD, COR_ARQ, COR_SPEC, COR_PLANO, "44546A", COR_VALIDAR, "1A5276"]
    for i, (row, cor) in enumerate(zip(rows, cores_fase)):
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            if j == 0:
                set_cell_bg(cell, cor)
                r = cell.paragraphs[0].add_run(val)
                r.bold = True
                r.font.size = Pt(11)
                set_run_color(r, BRANCO)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
                set_cell_bg(cell, bg)
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(10)

    add_spacer(doc)


# ─────────────────────────────────────────────
# 4. OS DOIS TIPOS DE TRABALHO
# ─────────────────────────────────────────────
def add_tipos_de_trabalho(doc):
    add_heading(doc, "3. Os Tipos de Trabalho", 1)

    doc.add_paragraph(
        "Todo trabalho no fluxo SDD se enquadra em uma de três categorias. "
        "Identificar o tipo correto é o primeiro passo antes de qualquer ação."
    )

    # Diagrama de decisão usando tabela
    add_spacer(doc)
    add_heading(doc, "Fluxo de Decisão", 3, COR_SUBTIT)

    decision_data = [
        # (cor_bg, texto, cor_texto)
        ("FDEBD0", "🤔  O que precisa ser feito?", PRETO),
        (None, None, None),  # spacer row
    ]

    # Tabela de decisão — 3 colunas para os 3 caminhos
    t = doc.add_table(rows=5, cols=3)
    set_table_no_spacing(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    labels_tipo = [
        (COR_PRD,     "🆕  NOVA FEATURE"),
        (COR_CR,      "🔧  ALTERAÇÃO / CORREÇÃO"),
        (COR_VALIDAR, "🐛  BUG SIMPLES"),
    ]
    passos = [
        ["PRD", "Change Request (CR)", "Change Request (CR)"],
        ["Arquitetura", "Avaliar impacto nos docs", "Implementar"],
        ["Spec Técnica", "Atualizar docs afetados", "Atualizar testes"],
        ["Plano → Código", "Implementar → Testar", "—"],
    ]

    # Linha 0: cabeçalhos dos tipos
    for j, (cor, label) in enumerate(labels_tipo):
        cell = t.cell(0, j)
        set_cell_bg(cell, cor)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        set_run_color(r, BRANCO)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    # Linhas de passos
    cores_linha = ["EBF5FB", "D6EAF8", "AED6F1", "85C1E9"]
    for i, (passo_row, cor_bg) in enumerate(zip(passos, cores_linha)):
        for j, passo in enumerate(passo_row):
            cell = t.cell(i + 1, j)
            set_cell_bg(cell, cor_bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arrow = "↓  " if i < len(passos) - 1 else "✅  "
            r = p.add_run(f"{arrow}{passo}" if passo != "—" else "—")
            r.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)

    # Regra importante
    add_info_box(
        doc,
        "⚠️  IMPORTANTE: O Change Request (CR) é OBRIGATÓRIO.\n"
        "Nunca implemente uma feature ou alteração significativa sem criar o CR primeiro.\n"
        "Mesmo para mudanças urgentes ou aparentemente simples.",
        "FBEEE6",
        bold=False
    )
    add_spacer(doc)


# ─────────────────────────────────────────────
# 5. GUIA DOS TEMPLATES
# ─────────────────────────────────────────────
def add_template_cr(doc):
    add_heading(doc, "4.1  Change Request (CR)", 2, COR_CR)

    # Cabeçalho colorido
    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_CR)
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/00-template-change-request.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, "85C1E9")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)

    add_heading(doc, "Propósito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "O Change Request é o ponto de entrada para qualquer alteração no projeto. "
        "Ele documenta O QUE vai mudar, POR QUE vai mudar, o impacto nos demais documentos "
        "e no código, as tarefas de implementação e os critérios de aceite."
    )

    add_heading(doc, "Quando criar", 3, COR_SUBTIT)
    items = [
        "Qualquer alteração em funcionalidade existente",
        "Correção de bug não trivial (que afeta lógica ou banco de dados)",
        "Nova feature significativa (antes de criar o PRD/Spec)",
        "Refatoração com impacto em múltiplos arquivos",
        "Mudança de dependência ou configuração de ambiente",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "Seções do Template", 3, COR_SUBTIT)

    secoes_cr = [
        ("1. Resumo da Mudança",
         "Descrição clara em 2–3 linhas: o que muda e por quê."),
        ("2. Classificação",
         "Tipo (Bug Fix / Nova Feature / Refactoring...), urgência e complexidade."),
        ("3. Contexto e Motivação",
         "AS-IS (situação atual), problema identificado, TO-BE (situação desejada)."),
        ("4. Detalhamento da Mudança",
         "Tabela AS-IS → TO-BE para cada regra/comportamento afetado. "
         "Também lista explicitamente o que NÃO muda."),
        ("5. Impacto nos Documentos",
         "Quais documentos (PRD, Arquitetura, Spec, Plano, Deploy Guide) são afetados "
         "e quais seções precisam ser atualizadas."),
        ("6. Impacto no Código",
         "Lista de arquivos a modificar/criar e descrição de cada migration necessária."),
        ("7. Tarefas de Implementação",
         "Tabela com ID (CR-T-01, CR-T-02...), descrição, dependências e critério Done When."),
        ("8. Critérios de Aceite",
         "Checklist de comportamentos esperados após a mudança."),
        ("9. Riscos e Efeitos Colaterais",
         "Riscos identificados, probabilidade, impacto e plano de mitigação."),
        ("10. Plano de Rollback",
         "Como reverter o código (git revert), a migration (alembic downgrade) "
         "e as variáveis de ambiente, se necessário."),
    ]

    t = doc.add_table(rows=len(secoes_cr), cols=2)
    t.style = "Table Grid"
    for i, (secao, desc) in enumerate(secoes_cr):
        c0 = t.cell(i, 0)
        c1 = t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg)
        set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(secao)
        r0.bold = True
        r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(desc)
        r1.font.size = Pt(10)

    add_spacer(doc)
    add_heading(doc, "Exemplo de uso", 3, COR_SUBTIT)
    doc.add_paragraph(
        "O desenvolvedor percebe que a funcionalidade de listagem de despesas precisa mostrar "
        "totais por status (Pago, Pendente, Atrasado). Antes de tocar no código:"
    )
    passos_ex = [
        "Cria docs/changes/CR-004-totais-por-status.md",
        "Preenche AS-IS: a listagem mostra apenas o total geral",
        "Preenche TO-BE: a listagem mostra total geral + subtotais por status",
        "Identifica que 03-SPEC.md, 04-IMPLEMENTATION-PLAN.md e o frontend precisam ser atualizados",
        "Lista as tarefas: atualizar backend (serviço + endpoint), atualizar frontend (componente de tabela)",
        "Só então começa a implementação, tarefa por tarefa",
    ]
    for i, passo in enumerate(passos_ex, 1):
        p = doc.add_paragraph(f"{i}. {passo}", style="List Number")
        p.runs[0].font.size = Pt(10)

    add_spacer(doc)


def add_template_prd(doc):
    add_heading(doc, "4.2  PRD — Product Requirements Document", 2, COR_PRD)

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_PRD)
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/01-template-prd.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, BRANCO)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_heading(doc, "Propósito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "O PRD define O QUÊ construir — os requisitos do produto do ponto de vista do usuário "
        "e do negócio. Não descreve como implementar; isso fica para a Arquitetura e a Spec. "
        "É o documento que justifica a existência de cada feature."
    )

    add_heading(doc, "Quando criar", 3, COR_SUBTIT)
    items = [
        "Início de um novo projeto",
        "Adição de um novo módulo grande (ex: módulo de pagamentos, módulo de relatórios)",
        "Quando um CR identifica impacto no PRD (novas user stories ou requisitos)",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "Seções do Template", 3, COR_SUBTIT)
    secoes = [
        ("1. Visão Geral",       "Descrição do produto, problema que resolve e público-alvo"),
        ("2. Objetivos e Métricas", "KPIs e metas mensuráveis para cada objetivo"),
        ("3. Personas",          "Perfis dos usuários principais com necessidades e frustrações"),
        ("4. Requisitos Funcionais", "Tabela RF-XXX com prioridade e detalhamento por bullet points"),
        ("5. Requisitos Não-Funcionais", "Performance, segurança, acessibilidade (RNF-XXX)"),
        ("6. User Stories",      "US-XXX: Como [persona], quero [ação], para que [benefício]"),
        ("7. Regras de Negócio", "RN-XXX: restrições e regras que o sistema deve respeitar"),
        ("8. Fora de Escopo",    "Listar explicitamente o que NÃO será construído nesta fase"),
        ("9. Dependências",      "Sistemas externos, APIs de terceiros, premissas assumidas"),
        ("10. Glossário",        "Definição dos termos do domínio de negócio"),
        ("Apêndice: Roadmap",    "Funcionalidades planejadas para fases futuras"),
    ]

    t = doc.add_table(rows=len(secoes), cols=2)
    t.style = "Table Grid"
    for i, (s, d) in enumerate(secoes):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(s); r0.bold = True; r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(d); r1.font.size = Pt(10)

    add_spacer(doc)
    add_info_box(
        doc,
        "💡  Dica: Mantenha os IDs dos requisitos (RF-001, RN-001, US-001) consistentes entre "
        "todos os documentos. A Spec vai referenciar RF-XXX; o Plano vai referenciar RF-XXX nas tarefas. "
        "Isso cria rastreabilidade completa.",
        COR_DESTAQUE
    )
    add_spacer(doc)


def add_template_arquitetura(doc):
    add_heading(doc, "4.3  Arquitetura", 2, COR_ARQ)

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_ARQ)
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/02-template-architecture.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, BRANCO)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_heading(doc, "Propósito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "A Arquitetura define COMO o sistema é construído: stack tecnológica, estrutura de pastas, "
        "modelagem de dados, padrões e convenções, estratégia de deploy e decisões técnicas. "
        "É o documento de referência para qualquer dúvida sobre \"como devemos fazer isso?\"."
    )

    add_heading(doc, "Quando criar / atualizar", 3, COR_SUBTIT)
    items = [
        "Início do projeto (definindo toda a stack)",
        "Ao adicionar uma nova tecnologia ou biblioteca importante",
        "Quando um CR introduz uma nova entidade de banco de dados (atualizar ER diagram)",
        "Ao estabelecer um novo padrão ou convenção que deve ser seguido pelo time",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "Seções do Template", 3, COR_SUBTIT)
    secoes = [
        ("1. Stack Tecnológica",        "Tabela: camada / tecnologia / versão / justificativa"),
        ("2. Arquitetura Geral",         "Diagrama Mermaid do fluxo entre componentes"),
        ("3. Estrutura de Pastas",       "Árvore completa de diretórios do projeto"),
        ("4. Modelagem de Dados",        "Diagrama ER em Mermaid + detalhamento de cada entidade"),
        ("5. Padrões e Convenções",      "Nomenclatura, Git (branches/commits), regras de API"),
        ("6. Integrações Externas",      "Serviços de terceiros com autenticação e links de docs"),
        ("7. Estratégia de Testes",      "Tipos de teste, ferramentas e cobertura mínima"),
        ("8. ADRs",                      "Registro de decisões técnicas com contexto e alternativas descartadas"),
        ("9. Deploy e Infraestrutura",   "Pipeline CI/CD, health check, rollback"),
        ("10. Gestão de Dependências",   "Política de pinning de versões e processo de auditoria"),
    ]

    t = doc.add_table(rows=len(secoes), cols=2)
    t.style = "Table Grid"
    for i, (s, d) in enumerate(secoes):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(s); r0.bold = True; r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(d); r1.font.size = Pt(10)

    add_spacer(doc)
    add_info_box(
        doc,
        "📋  ADRs (Architecture Decision Records): Sempre que você toma uma decisão técnica "
        "não óbvia (ex: \"Por que usamos SQLite em dev em vez de PostgreSQL?\"), registre como "
        "ADR-XXX na seção 8. Inclua as alternativas que foram descartadas e por quê. "
        "Isso evita que a equipe reabra discussões já resolvidas.",
        COR_DESTAQUE
    )
    add_spacer(doc)


def add_template_spec(doc):
    add_heading(doc, "4.4  Especificação Técnica", 2, COR_SPEC)

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_SPEC)
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/03-template-spec.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, BRANCO)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_heading(doc, "Propósito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "A Spec é o documento mais técnico do fluxo. Ela detalha EXATAMENTE o que o desenvolvedor "
        "(ou agente de IA) deve implementar: contratos de API, tipos/interfaces, lógica de negócio "
        "passo a passo, schema de banco de dados, validações, casos de borda e plano de testes. "
        "Não deve haver ambiguidade — se algo não está na Spec, não está no escopo."
    )

    add_heading(doc, "Quando criar / atualizar", 3, COR_SUBTIT)
    items = [
        "Antes de implementar qualquer nova feature",
        "Quando um CR altera contratos de API ou lógica de negócio existente",
        "Para documentar retroativamente features que foram implementadas sem spec",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "Seções do Template", 3, COR_SUBTIT)
    secoes = [
        ("1. Resumo das Mudanças",           "Visão geral e escopo da iteração"),
        ("Contratos da API (Visão Geral)",    "Tabela macro de todos os endpoints: método, path, auth, body, resposta"),
        ("2. Detalhamento Técnico por Feature", "Para cada RF-XXX: arquivos, interfaces/types, lógica de negócio, endpoints, banco de dados, validações"),
        ("3. Componentes de UI",              "Props, estados (loading, empty, error) e comportamentos dos componentes"),
        ("4. Fluxos Críticos",               "Diagramas sequenciais Mermaid dos fluxos principais"),
        ("5. Casos de Borda",                 "Tabela de cenários não óbvios e comportamento esperado"),
        ("6. Plano de Testes",               "IDs de teste (BT-XXX para backend, FT-XXX para fluxo completo)"),
        ("7. Checklist de Implementação",    "Lista de verificação para confirmar completude"),
    ]

    t = doc.add_table(rows=len(secoes), cols=2)
    t.style = "Table Grid"
    for i, (s, d) in enumerate(secoes):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(s); r0.bold = True; r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(d); r1.font.size = Pt(10)

    add_spacer(doc)
    add_info_box(
        doc,
        "💡  Dica — Tabela de Contratos da API: Antes de detalhar cada feature, inclua uma tabela "
        "resumo com todos os endpoints do módulo (método, path, autenticação, corpo, resposta). "
        "Isso dá uma visão macro para o desenvolvedor antes de mergulhar nos detalhes.\n\n"
        "💡  Padrão PATCH: Use sempre PATCH (não PUT) para atualizações parciais. "
        "No backend, aplique exclude_unset=True (Pydantic) para ignorar campos não enviados.",
        COR_DESTAQUE
    )
    add_spacer(doc)


def add_template_plano(doc):
    add_heading(doc, "4.5  Plano de Implementação", 2, COR_PLANO)

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_PLANO)
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/04-template-implementation-plan.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, BRANCO)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_heading(doc, "Propósito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "O Plano de Implementação traduz os requisitos da Spec em tarefas concretas e sequenciadas. "
        "Define a ordem de execução, as dependências entre tarefas e o critério de conclusão (Done When) "
        "para cada uma. É o roadmap técnico do desenvolvedor."
    )

    add_heading(doc, "Estrutura de Tarefas", 3, COR_SUBTIT)
    doc.add_paragraph(
        "As tarefas são organizadas em grupos (1–6) e cada tarefa tem um ID único (T-XXX). "
        "CRs adicionados depois do plano inicial recebem seções próprias com IDs CR-T-XX."
    )

    grupos = [
        ("Grupo 1", "Setup e Infraestrutura",  "T-001 a T-005 — configuração inicial do projeto"),
        ("Grupo 2", "Modelos e Entidades",      "T-006 a T-008 — migrations, models, types"),
        ("Grupo 3", "Lógica de Negócio",        "T-009 a T-013 — services e validações"),
        ("Grupo 4", "API / Controllers",        "T-014 a T-021 — endpoints e middlewares"),
        ("Grupo 5", "UI / Frontend",            "T-022 a T-028 — componentes e páginas"),
        ("Grupo 6", "Testes e Refinamento",     "T-029 a T-033 — testes de integração e E2E"),
    ]

    t = doc.add_table(rows=len(grupos), cols=3)
    t.style = "Table Grid"
    for i, (grupo, nome, desc) in enumerate(grupos):
        c0, c1, c2 = t.cell(i, 0), t.cell(i, 1), t.cell(i, 2)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, COR_PLANO); set_cell_bg(c1, bg); set_cell_bg(c2, bg)
        r0 = c0.paragraphs[0].add_run(grupo); r0.bold = True; r0.font.size = Pt(10); set_run_color(r0, BRANCO)
        r1 = c1.paragraphs[0].add_run(nome); r1.bold = True; r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(desc); r2.font.size = Pt(10)

    add_spacer(doc)
    add_heading(doc, "Coluna Done When", 3, COR_SUBTIT)
    doc.add_paragraph(
        "Cada tarefa tem um critério objetivo de conclusão na coluna \"Done When\". "
        "Exemplos válidos: \"Migration roda sem erro\", \"Endpoint retorna 201 com body correto\", "
        "\"Componente renderiza sem erros no console\". Isso elimina subjetividade sobre "
        "o que significa \"terminar\" uma tarefa."
    )

    add_spacer(doc)
    add_info_box(
        doc,
        "📋  CRs no Plano: Quando um CR é aprovado, adicione uma nova seção ## CR-XXX ao final "
        "do arquivo 04-IMPLEMENTATION-PLAN.md. Use a mesma estrutura de grupos e tarefas, "
        "mas com IDs CR-T-01, CR-T-02, etc. Isso mantém todo o histórico de implementação "
        "em um único documento.",
        COR_DESTAQUE
    )
    add_spacer(doc)


def add_template_claude(doc):
    add_heading(doc, "4.6  CLAUDE.md — Instruções para o Agente de IA", 2, "5D3FD3")

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "5D3FD3")
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/CLAUDE-template.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, BRANCO)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_heading(doc, "Propósito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "O CLAUDE.md é o arquivo de instruções persistentes para agentes de IA (como o Claude). "
        "Ele define a identidade do projeto, o fluxo de trabalho que deve ser seguido, "
        "as regras de implementação, a stack tecnológica, o contexto atual e o histórico "
        "de problemas conhecidos. "
        "Toda nova conversa com o agente começa com este arquivo carregado automaticamente."
    )

    add_heading(doc, "O que incluir", 3, COR_SUBTIT)
    items = [
        ("Identidade do Projeto",          "Nome, descrição resumida, stack e URL do repositório"),
        ("Fluxo de Desenvolvimento",        "Tabela de fases + Regra de Ouro + instrução de CR obrigatório"),
        ("Templates de Documentos",         "Tabela: tipo de doc → caminho do template"),
        ("Regras de Implementação",         "Antes de Codar (5 regras), Done When Universal, formato de commits"),
        ("Regras para Alterações",          "Fluxo CR em 8 passos"),
        ("Estrutura de Pastas",             "Árvore do projeto (sincronizada com o código real)"),
        ("Convenções de Código",            "Nomenclatura por tipo de arquivo e função"),
        ("Stack Tecnológica",               "Tabela completa com versões"),
        ("Contexto Atual",                  "CRs ativos, última tarefa implementada, documentos existentes"),
        ("Lembretes Importantes",           "Regras de comportamento do agente (não inventar, um passo por vez, etc.)"),
        ("Troubleshooting",                 "Problemas já resolvidos: dependências, banco de dados, ambiente"),
    ]

    t = doc.add_table(rows=len(items), cols=2)
    t.style = "Table Grid"
    for i, (s, d) in enumerate(items):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(s); r0.bold = True; r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(d); r1.font.size = Pt(10)

    add_spacer(doc)
    add_info_box(
        doc,
        "🔄  Mantenha o CLAUDE.md sempre atualizado. Ao concluir cada CR:\n"
        "1. Atualize a seção \"Contexto Atual\" (CRs ativos, última tarefa)\n"
        "2. Atualize a estrutura de pastas se novos arquivos foram criados\n"
        "3. Adicione à seção Troubleshooting qualquer problema novo que foi resolvido\n\n"
        "Um CLAUDE.md desatualizado faz o agente tomar decisões baseadas em contexto errado.",
        COR_DESTAQUE
    )
    add_spacer(doc)


def add_guia_templates(doc):
    add_heading(doc, "4. Guia de Cada Template", 1)
    doc.add_paragraph(
        "Esta seção apresenta o propósito, quando usar e as seções de cada template "
        "disponível em docs/templates/. Use sempre o template correspondente ao criar "
        "qualquer documento do fluxo."
    )
    add_spacer(doc)
    add_template_cr(doc)
    add_template_prd(doc)
    add_template_arquitetura(doc)
    add_template_spec(doc)
    add_template_plano(doc)
    add_template_claude(doc)


# ─────────────────────────────────────────────
# 6. SISTEMA DE IDENTIFICADORES
# ─────────────────────────────────────────────
def add_identificadores(doc):
    add_heading(doc, "5. Sistema de Identificadores", 1)
    doc.add_paragraph(
        "O fluxo SDD usa um sistema de IDs para criar rastreabilidade entre documentos. "
        "Um ID em um documento deve aparecer exatamente igual em todos os outros documentos "
        "que fazem referência a ele."
    )

    ids = [
        ("RF-XXX",  COR_PRD,     "Requisito Funcional",           "PRD → Spec → Plano",               "RF-013: Módulo de Gastos Diários"),
        ("RN-XXX",  COR_ARQ,     "Regra de Negócio",              "PRD → Spec",                       "RN-019: Gastos diários não participam da transição de mês"),
        ("US-XXX",  COR_SPEC,    "User Story",                    "PRD",                              "US-020: Como usuário, quero registrar um gasto diário"),
        ("T-XXX",   COR_PLANO,   "Tarefa de Implementação",       "Plano de Implementação",           "T-045: Criar modelo DailyExpense"),
        ("CR-XXX",  COR_CR,      "Change Request",                "Todas as fases",                   "CR-005: Gastos Diários"),
        ("CR-T-XX", COR_CR,      "Tarefa de um CR específico",    "Plano (seção do CR)",              "CR-T-01: Criar migration daily_expenses"),
        ("ADR-XXX", COR_ARQ,     "Architecture Decision Record",  "Arquitetura",                      "ADR-014: Uso de Alembic para migrations"),
        ("BT-XXX",  COR_VALIDAR, "Backend Test",                  "Spec → Testes",                    "BT-011: POST /api/daily-expenses retorna 201"),
        ("FT-XXX",  COR_VALIDAR, "Full Flow Test",                "Spec → Testes",                    "FT-030: Usuário registra gasto e vê na listagem"),
    ]

    t = doc.add_table(rows=1 + len(ids), cols=5)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Prefixo", "Tipo", "Usado em", "Exemplo"]
    for j, h in enumerate([headers[0]] + headers[1:]):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        r = cell.paragraphs[0].add_run(["Prefixo", "Tipo", "Cor", "Usado em", "Exemplo"][j])
        r.bold = True; r.font.size = Pt(10); set_run_color(r, BRANCO)

    for i, (prefixo, cor, tipo, usado_em, exemplo) in enumerate(ids):
        row_data = [prefixo, tipo, usado_em, exemplo]
        for j, val in enumerate(row_data):
            cell = t.cell(i + 1, j)
            if j == 0:
                set_cell_bg(cell, cor)
                r = cell.paragraphs[0].add_run(val)
                r.bold = True; r.font.size = Pt(10); set_run_color(r, BRANCO)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
                set_cell_bg(cell, bg)
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(9 if j == 3 else 10)

    add_spacer(doc)
    add_info_box(
        doc,
        "🔗  Rastreabilidade: Se o RF-013 aparecer no PRD, ele deve aparecer como referência "
        "na Spec (seção da feature) e no Plano (coluna Ref das tarefas). "
        "Isso permite rastrear qualquer requisito de ponta a ponta.",
        COR_DESTAQUE
    )
    add_spacer(doc)


# ─────────────────────────────────────────────
# 7. PROMPT DE INÍCIO DE PROJETO
# ─────────────────────────────────────────────
def add_prompt_inicio(doc):
    add_heading(doc, "6. Prompt para Iniciar Novo Projeto", 1)

    doc.add_paragraph(
        "Use o seguinte prompt ao iniciar uma nova conversa com um agente de IA (como o Claude) "
        "para configurar o fluxo SDD em um novo projeto. Adapte os campos entre colchetes."
    )

    add_spacer(doc)
    add_heading(doc, "Prompt de Setup Inicial", 3, COR_SUBTIT)

    prompt_texto = """\
Vamos iniciar um novo projeto de software chamado [NOME DO PROJETO].

## Contexto do Projeto
- **Descrição:** [O que o sistema faz e para quem]
- **Stack:** [Frontend: ex. React + TypeScript] / [Backend: ex. FastAPI + Python] / [Banco: ex. PostgreSQL]
- **Repositório:** [URL do Git ou "criar novo"]

## O que preciso que você faça

**Passo 1 — Criar o CLAUDE.md:**
Leia o template em docs/templates/CLAUDE-template.md e crie o arquivo CLAUDE.md
na raiz do projeto. Substitua todos os [placeholders] pelas informações acima.

**Passo 2 — Criar o PRD (01-PRD.md):**
Leia o template em docs/templates/01-template-prd.md e crie docs/01-PRD.md
com base nas informações abaixo:
[DESCREVER AS FUNCIONALIDADES PRINCIPAIS DO PRODUTO]

**Passo 3 — Criar a Arquitetura (02-ARCHITECTURE.md):**
Com base no PRD, crie docs/02-ARCHITECTURE.md usando o template
docs/templates/02-template-architecture.md.

**Passo 4 — Criar a Spec Técnica (03-SPEC.md):**
Com base no PRD e na Arquitetura, crie docs/03-SPEC.md usando o template
docs/templates/03-template-spec.md. Detalhe os endpoints, tipos, lógica
de negócio e plano de testes para cada requisito funcional.

**Passo 5 — Criar o Plano de Implementação (04-IMPLEMENTATION-PLAN.md):**
Com base na Spec, crie docs/04-IMPLEMENTATION-PLAN.md usando o template
docs/templates/04-template-implementation-plan.md.

## Regras que você deve seguir
- Documente PRIMEIRO, código DEPOIS
- Nunca invente funcionalidades que não estão no PRD
- Use os IDs RF-XXX, RN-XXX, T-XXX, CR-XXX conforme o fluxo SDD
- Ao terminar, confirme que todos os 5 documentos foram criados\
"""

    add_code_box(doc, prompt_texto)
    add_spacer(doc)

    add_heading(doc, "Prompt para Novo CR (Feature ou Correção)", 3, COR_SUBTIT)

    prompt_cr = """\
Preciso implementar uma nova alteração no projeto [NOME DO PROJETO].

## Descrição da Mudança
[DESCREVER O QUE PRECISA SER FEITO]

## Instruções
1. Leia o arquivo CLAUDE.md para entender o contexto atual do projeto
2. Crie um Change Request usando o template docs/templates/00-template-change-request.md
   - Salve em docs/changes/CR-[PRÓXIMO NÚMERO]-[slug].md
   - Preencha todas as seções: AS-IS, TO-BE, impacto nos docs, tarefas, critérios de aceite
3. Avalie o impacto e atualize os documentos afetados (PRD, Arquitetura, Spec, Plano)
4. Implemente as tarefas do CR uma por vez, confirmando cada Done When
5. Ao terminar: verifique o build, escreva os testes, atualize o CLAUDE.md, faça o commit\
"""

    add_code_box(doc, prompt_cr)
    add_spacer(doc)

    add_info_box(
        doc,
        "💡  Dica: Seja específico na descrição da mudança. Em vez de \"melhorar a listagem de "
        "despesas\", escreva \"adicionar totalizadores por status (Pago, Pendente, Atrasado) "
        "na tabela de despesas mensais, abaixo do total geral\". Quanto mais específico, "
        "menos ambiguidade para o agente resolver.",
        COR_DESTAQUE
    )
    add_spacer(doc)


# ─────────────────────────────────────────────
# 8. DONE WHEN UNIVERSAL
# ─────────────────────────────────────────────
def add_done_when(doc):
    add_heading(doc, "7. Checklist \"Done When\" Universal", 1)

    doc.add_paragraph(
        "Toda tarefa (T-XXX ou CR-T-XX) só é considerada concluída quando todos os critérios "
        "obrigatórios abaixo são satisfeitos. Use esta checklist como revisão antes de fazer o commit."
    )

    add_spacer(doc)

    # Obrigatórios
    add_heading(doc, "Obrigatórios (todas as tarefas)", 3, COR_SUBTIT)
    obrigatorios = [
        "Funcionalidade implementada conforme descrito na tarefa",
        "App roda localmente sem erros (backend + frontend)",
        "Testes existentes continuam passando (sem regressão)",
        "Novos testes cobrem a funcionalidade adicionada ou alterada",
        "Commit segue Conventional Commits e referencia o ID da tarefa",
    ]
    t_obrig = doc.add_table(rows=len(obrigatorios), cols=2)
    for i, item in enumerate(obrigatorios):
        c0, c1 = t_obrig.cell(i, 0), t_obrig.cell(i, 1)
        set_cell_bg(c0, "D5F5E3")
        r0 = c0.paragraphs[0].add_run("✅")
        r0.font.size = Pt(12)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c1, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r1 = c1.paragraphs[0].add_run(item)
        r1.font.size = Pt(11)

    add_spacer(doc)

    # Se aplicável
    add_heading(doc, "Se aplicável", 3, COR_SUBTIT)
    aplicaveis = [
        ("Migration",     "alembic upgrade head + alembic downgrade -1 testados sem erros"),
        ("Endpoints",     "Respondem com status codes corretos (201, 204, 400, 401, 404, etc.)"),
        ("Documentação",  "Documentos afetados atualizados (Spec, Arquitetura, CLAUDE.md)"),
        ("Frontend",      "Sem erros ou warnings no console do browser"),
        ("Build",         "Build TypeScript passa sem erros de tipo (npx tsc --noEmit)"),
    ]

    t_aplic = doc.add_table(rows=len(aplicaveis), cols=3)
    t_aplic.style = "Table Grid"
    for i, (contexto, item) in enumerate(aplicaveis):
        c0, c1, c2 = t_aplic.cell(i, 0), t_aplic.cell(i, 1), t_aplic.cell(i, 2)
        set_cell_bg(c0, "FEF9E7")
        set_cell_bg(c1, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        set_cell_bg(c2, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r0 = c0.paragraphs[0].add_run("Se aplicável"); r0.font.size = Pt(9); r0.italic = True
        r1 = c1.paragraphs[0].add_run(contexto); r1.bold = True; r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(item); r2.font.size = Pt(10)

    add_spacer(doc)


# ─────────────────────────────────────────────
# 9. REFERÊNCIA RÁPIDA
# ─────────────────────────────────────────────
def add_referencia_rapida(doc):
    add_heading(doc, "8. Referência Rápida", 1)

    # Tabela: Situação → Ação
    add_heading(doc, "Quando criar cada documento", 3, COR_SUBTIT)

    situacoes = [
        ("Vou alterar algo que já existe",    "Change Request (CR)",           "00-template-change-request.md", COR_CR),
        ("Vou criar um módulo novo",          "PRD + Arquitetura + Spec + Plano", "01 a 04 templates",          COR_PRD),
        ("Vou mudar a stack ou estrutura",    "Arquitetura + CR",              "02-template-architecture.md",   COR_ARQ),
        ("Vou implementar uma nova feature",  "Spec + Plano (ou CR se existente)", "03-template-spec.md",       COR_SPEC),
        ("Vou iniciar um novo projeto",       "CLAUDE.md + todos os documentos", "CLAUDE-template.md + 01-04",  "5D3FD3"),
        ("Vou fazer deploy",                  "Deploy Guide",                  "05-DEPLOY-GUIDE.md",           "1A5276"),
    ]

    t = doc.add_table(rows=1 + len(situacoes), cols=3)
    t.style = "Table Grid"
    headers_r = ["Situação", "Documento a Criar", "Template"]
    for j, h in enumerate(headers_r):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10); set_run_color(r, BRANCO)

    for i, (sit, doc_criar, template, cor) in enumerate(situacoes):
        for j, val in enumerate([sit, doc_criar, template]):
            cell = t.cell(i + 1, j)
            bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
            if j == 1:
                set_cell_bg(cell, cor)
                r = cell.paragraphs[0].add_run(val)
                r.bold = True; r.font.size = Pt(10); set_run_color(r, BRANCO)
            else:
                set_cell_bg(cell, bg)
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(10)

    add_spacer(doc)

    # Conventional Commits
    add_heading(doc, "Conventional Commits", 3, COR_SUBTIT)

    commits = [
        ("feat:",     "Nova feature implementada",           "feat: implement CR-005 - daily expenses module"),
        ("fix:",      "Correção de bug",                    "fix: CR-007 - prevent duplicate installments on month transition"),
        ("docs:",     "Atualização de documentação",        "docs: update spec and PRD for CR-005"),
        ("refactor:", "Refatoração sem mudança de comportamento", "refactor: extract status logic to dedicated service"),
        ("test:",     "Adição ou correção de testes",       "test: add BT-011 to BT-015 for daily expenses endpoints"),
        ("chore:",    "Tarefas de manutenção",              "chore: update dependencies, pin bcrypt==4.0.*"),
    ]

    t2 = doc.add_table(rows=len(commits), cols=3)
    t2.style = "Table Grid"
    for i, (prefixo, desc, exemplo) in enumerate(commits):
        c0, c1, c2 = t2.cell(i, 0), t2.cell(i, 1), t2.cell(i, 2)
        set_cell_bg(c0, COR_TITULO)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c1, bg); set_cell_bg(c2, bg)
        r0 = c0.paragraphs[0].add_run(prefixo); r0.font.name = "Courier New"; r0.font.size = Pt(10); set_run_color(r0, "85C1E9")
        r1 = c1.paragraphs[0].add_run(desc); r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(exemplo); r2.font.name = "Courier New"; r2.font.size = Pt(9)

    add_spacer(doc)

    # Fluxo CR em 8 passos
    add_heading(doc, "Fluxo CR em 8 Passos", 3, COR_SUBTIT)

    passos_cr = [
        "Criar o CR em docs/changes/CR-[XXX]-[slug].md usando o template",
        "Preencher todas as seções: resumo, AS-IS/TO-BE, impacto, tarefas, critérios de aceite",
        "Avaliar impacto nos documentos: quais seções de PRD, Spec, Arquitetura, Plano são afetados?",
        "Atualizar os documentos afetados ANTES de implementar",
        "Implementar as tarefas do CR uma por vez",
        "Validar os critérios de aceite do CR + checklist Done When Universal",
        "Verificar o build: npx tsc --noEmit (frontend) + testes passando (backend)",
        "Commit e push referenciando o CR: feat: CR-XXX - descrição",
    ]

    t3 = doc.add_table(rows=len(passos_cr), cols=2)
    cores_passo = [COR_CR, COR_CR, COR_PRD, COR_ARQ, COR_SPEC, COR_PLANO, COR_VALIDAR, "1A5276"]
    for i, (passo, cor) in enumerate(zip(passos_cr, cores_passo)):
        c0, c1 = t3.cell(i, 0), t3.cell(i, 1)
        set_cell_bg(c0, cor)
        set_cell_bg(c1, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r0 = c0.paragraphs[0].add_run(str(i + 1))
        r0.bold = True; r0.font.size = Pt(14); set_run_color(r0, BRANCO)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = c1.paragraphs[0].add_run(passo); r1.font.size = Pt(11)

    add_spacer(doc, 2)
    add_info_box(
        doc,
        "📄  Este manual está disponível em: docs/manual-workflow-sdd.docx\n"
        "Os templates estão em: docs/templates/\n"
        "O template genérico do CLAUDE.md está em: docs/templates/CLAUDE-template.md",
        COR_CINZA_BG
    )
    add_spacer(doc)


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def main():
    doc = Document()
    configure_document(doc)

    add_cover(doc)
    add_introducao(doc)
    add_visao_geral(doc)
    add_page_break(doc)
    add_tipos_de_trabalho(doc)
    add_page_break(doc)
    add_guia_templates(doc)
    add_page_break(doc)
    add_identificadores(doc)
    add_page_break(doc)
    add_prompt_inicio(doc)
    add_page_break(doc)
    add_done_when(doc)
    add_referencia_rapida(doc)

    # Determinar caminho de saída
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, "manual-workflow-sdd.docx")
    doc.save(output_path)
    print(f"[OK]  Manual gerado: {output_path}")


if __name__ == "__main__":
    main()
