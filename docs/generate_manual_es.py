"""
Script para generar el Manual del Flujo de Trabajo Spec-Driven Development (version en espanol uruguayo).
Salida: docs/manual-workflow-sdd-es.docx

Uso: python docs/generate_manual_es.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import date

# ─────────────────────────────────────────────
# Paleta de colores
# ─────────────────────────────────────────────
COR_CR       = "1B4F72"
COR_PRD      = "2471A3"
COR_ARQ      = "145A32"
COR_SPEC     = "1E8449"
COR_PLANO    = "B7950B"
COR_CODIGO   = "566573"
COR_VALIDAR  = "922B21"
COR_TITULO   = "1A1A2E"
COR_SUBTIT   = "16213E"
COR_DESTAQUE = "E8F4F8"
COR_CINZA_BG = "F2F3F4"
BRANCO       = "FFFFFF"
PRETO        = "000000"


# ─────────────────────────────────────────────
# Helpers de XML / formatacao
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_run_color(run, hex_color: str):
    run.font.color.rgb = RGBColor.from_string(hex_color)


def set_table_no_spacing(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    cellMar = OxmlElement("w:tblCellMar")
    for side in ["top", "bottom", "left", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "80")
        el.set(qn("w:type"), "dxa")
        cellMar.append(el)
    tblPr.append(cellMar)


def add_page_break(doc):
    doc.add_page_break()


def add_spacer(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_heading(doc, text, level=1, color=COR_TITULO):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_color(run, color)
    return h


def add_info_box(doc, text, bg_color=COR_CINZA_BG, bold=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return table


def add_code_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1E1E2E")
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    set_run_color(run, "A8FF60")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return table


# ─────────────────────────────────────────────
# Estilos globales del documento
# ─────────────────────────────────────────────
def configure_document(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i, sz in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {i}"]
        h.font.name = "Calibri"
        h.font.size = Pt(sz)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after  = Pt(6)


# ─────────────────────────────────────────────
# 1. PORTADA
# ─────────────────────────────────────────────
def add_portada(doc):
    add_spacer(doc, 5)

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_TITULO)
    cell.width = Inches(6.5)

    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("Guia del Flujo de Trabajo")
    r.font.name = "Calibri"
    r.font.size = Pt(26)
    r.bold = True
    set_run_color(r, BRANCO)
    p_title.paragraph_format.space_before = Pt(16)

    p_sub = cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("Spec-Driven Development")
    r2.font.name = "Calibri"
    r2.font.size = Pt(20)
    r2.bold = True
    set_run_color(r2, "85C1E9")

    p_desc = cell.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_desc.add_run("Un metodo para desarrollar software con documentacion como base")
    r3.font.name = "Calibri"
    r3.font.size = Pt(12)
    r3.italic = True
    set_run_color(r3, "BDC3C7")
    p_desc.paragraph_format.space_after = Pt(16)

    add_spacer(doc, 2)

    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p_ver.add_run(f"Version 1.0  *  {date.today().strftime('%d/%m/%Y')}")
    r4.font.size = Pt(10)
    set_run_color(r4, "717D7E")

    add_page_break(doc)


# ─────────────────────────────────────────────
# 2. INTRODUCCION
# ─────────────────────────────────────────────
def add_introduccion(doc):
    add_heading(doc, "1. Introduccion", 1)

    doc.add_paragraph(
        "El Spec-Driven Development (SDD) es un enfoque de desarrollo de software en el que "
        "la documentacion se crea antes del codigo. Cada funcionalidad o cambio empieza con la "
        "definicion escrita de lo que hay que construir: el 'por que', el 'que' y el 'como', "
        "antes de que se escriba cualquier linea de codigo."
    )

    add_heading(doc, "?Por que documentar antes de codear?", 3, COR_SUBTIT)

    items = [
        ("Trazabilidad total",
         "Cada decision tiene origen documentado. Siempre sabe por que algo se hizo de esa manera."),
        ("Menos retrabajo",
         "Los malentendidos se resuelven en el papel, no en el codigo. Refactorizar un documento es "
         "infinitamente mas barato que refactorizar codigo en produccion."),
        ("Onboarding acelerado",
         "Los nuevos integrantes del equipo (incluidos los agentes de IA) entienden el proyecto "
         "leyendo los documentos, sin tener que rebuscar en el codigo fuente."),
        ("Alcance controlado",
         "La spec es el contrato. Si no esta en la spec, no esta en el alcance. Esto evita el "
         "'scope creep' y los pedidos de features no documentados."),
        ("Decisiones preservadas",
         "Los Architecture Decision Records (ADRs) registran las decisiones tecnicas y sus "
         "alternativas descartadas, evitando que el equipo reabra discusiones ya resueltas."),
    ]

    for titulo, descripcion in items:
        p = doc.add_paragraph(style="List Bullet")
        r_bold = p.add_run(f"{titulo}: ")
        r_bold.bold = True
        r_bold.font.size = Pt(11)
        p.add_run(descripcion).font.size = Pt(11)

    add_spacer(doc)
    add_info_box(
        doc,
        "Regla de Oro: Documentacion PRIMERO -> Codigo DESPUES.\n"
        "Si no podes escribir lo que vas a construir, todavia no entendiste bien el problema.",
        COR_DESTAQUE,
        bold=False
    )
    add_spacer(doc)


# ─────────────────────────────────────────────
# 3. VISION GENERAL DEL FLUJO
# ─────────────────────────────────────────────
def add_vision_general(doc):
    add_heading(doc, "2. Vision General del Flujo", 1)

    doc.add_paragraph(
        "El flujo SDD tiene 8 fases numeradas del 0 al 7. Cada fase tiene un documento "
        "correspondiente y un proposito especifico. El flujo completo va desde la definicion "
        "del problema hasta el deploy en produccion."
    )

    add_spacer(doc)

    fases = [
        (COR_CR,     "Fase 0\nCR"),
        (COR_PRD,    "Fase 1\nPRD"),
        (COR_ARQ,    "Fase 2\nArquit."),
        (COR_SPEC,   "Fase 3\nSpec"),
        (COR_PLANO,  "Fase 4\nPlan"),
        (COR_CODIGO, "Fase 5\nCodigo"),
        (COR_VALIDAR,"Fase 6\nValid."),
        ("1A5276",   "Fase 7\nDeploy"),
    ]

    table = doc.add_table(rows=2, cols=len(fases))
    set_table_no_spacing(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (cor, label) in enumerate(fases):
        cell = table.cell(0, i)
        set_cell_bg(cell, cor)
        cell.width = Inches(0.85)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.size = Pt(8.5)
        r.bold = True
        set_run_color(r, BRANCO)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    arrow_texts = ["->", "->", "->", "->", "->", "->", "->", ""]
    for i, arrow in enumerate(arrow_texts):
        cell = table.cell(1, i)
        set_cell_bg(cell, "FAFAFA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(arrow)
        r.font.size = Pt(14)
        set_run_color(r, "808B96")

    add_spacer(doc)

    add_heading(doc, "Referencia de las Fases", 3, COR_SUBTIT)

    rows = [
        ["0", "Change Request (CR)", "/docs/changes/CR-XXX.md",
         "Cambios y correcciones en funcionalidades existentes"],
        ["1", "PRD", "/docs/01-PRD.md",
         "Definicion inicial o incorporacion de modulos grandes"],
        ["2", "Arquitectura", "/docs/02-ARCHITECTURE.md",
         "Decisiones de stack, estructura y patrones"],
        ["3", "Spec Tecnica", "/docs/03-SPEC.md",
         "Detalle tecnico de cada funcionalidad"],
        ["4", "Plan de Implementacion", "/docs/04-IMPLEMENTATION-PLAN.md",
         "Orden y dependencias de las tareas"],
        ["5", "Implementacion", "Codigo fuente",
         "Construccion efectiva"],
        ["6", "Validacion", "Checklist 'Done When'",
         "Verificar criterios de aceptacion antes del deploy"],
        ["7", "Deploy y Release", "/docs/05-DEPLOY-GUIDE.md",
         "Procedimientos de deploy, rollback y verificacion"],
    ]

    t = doc.add_table(rows=1 + len(rows), cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Fase", "Documento", "Ruta", "Cuando usarlo"]
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        set_run_color(r, BRANCO)

    cores_fase = [COR_CR, COR_PRD, COR_ARQ, COR_SPEC, COR_PLANO, "44546A", COR_VALIDAR, "1A5276"]
    for i, (row, cor) in enumerate(zip(rows, cores_fase)):
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            if j == 0:
                set_cell_bg(cell, cor)
                r = cell.paragraphs[0].add_run(val)
                r.bold = True
                r.font.size = Pt(11)
                set_run_color(r, BRANCO)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
                set_cell_bg(cell, bg)
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(10)

    add_spacer(doc)


# ─────────────────────────────────────────────
# 4. LOS TIPOS DE TRABAJO
# ─────────────────────────────────────────────
def add_tipos_de_trabajo(doc):
    add_heading(doc, "3. Los Tipos de Trabajo", 1)

    doc.add_paragraph(
        "Todo trabajo en el flujo SDD corresponde a una de tres categorias. "
        "Identificar el tipo correcto es el primer paso antes de cualquier accion."
    )

    add_spacer(doc)
    add_heading(doc, "Flujo de Decision", 3, COR_SUBTIT)

    t = doc.add_table(rows=5, cols=3)
    set_table_no_spacing(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    labels_tipo = [
        (COR_PRD,     "NUEVA FUNCIONALIDAD"),
        (COR_CR,      "CAMBIO / CORRECCION"),
        (COR_VALIDAR, "BUG SIMPLE"),
    ]
    passos = [
        ["PRD", "Change Request (CR)", "Change Request (CR)"],
        ["Arquitectura", "Evaluar impacto en docs", "Implementar"],
        ["Spec Tecnica", "Actualizar docs afectados", "Actualizar tests"],
        ["Plan -> Codigo", "Implementar -> Testear", "--"],
    ]

    for j, (cor, label) in enumerate(labels_tipo):
        cell = t.cell(0, j)
        set_cell_bg(cell, cor)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        set_run_color(r, BRANCO)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    cores_linha = ["EBF5FB", "D6EAF8", "AED6F1", "85C1E9"]
    for i, (passo_row, cor_bg) in enumerate(zip(passos, cores_linha)):
        for j, passo in enumerate(passo_row):
            cell = t.cell(i + 1, j)
            set_cell_bg(cell, cor_bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            flecha = "v  " if i < len(passos) - 1 else "OK  "
            r = p.add_run(f"{flecha}{passo}" if passo != "--" else "--")
            r.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_info_box(
        doc,
        "IMPORTANTE: El Change Request (CR) es OBLIGATORIO.\n"
        "Nunca implementes una funcionalidad o cambio significativo sin crear el CR primero.\n"
        "Incluso para cambios urgentes o aparentemente simples.",
        "FBEEE6",
        bold=False
    )
    add_spacer(doc)


# ─────────────────────────────────────────────
# 5. GUIA DE TEMPLATES
# ─────────────────────────────────────────────
def add_template_cr(doc):
    add_heading(doc, "4.1  Change Request (CR)", 2, COR_CR)

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_CR)
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/00-template-change-request.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, "85C1E9")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_heading(doc, "Proposito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "El Change Request es el punto de entrada para cualquier cambio en el proyecto. "
        "Documenta QUE va a cambiar, POR QUE va a cambiar, el impacto en los demas documentos "
        "y en el codigo, las tareas de implementacion y los criterios de aceptacion."
    )

    add_heading(doc, "Cuando crear uno", 3, COR_SUBTIT)
    items = [
        "Cualquier cambio en funcionalidad existente",
        "Correccion de bug no trivial (que afecta logica o base de datos)",
        "Nueva funcionalidad significativa (antes de crear el PRD/Spec)",
        "Refactorizacion con impacto en multiples archivos",
        "Cambio de dependencia o configuracion de entorno",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "Secciones del Template", 3, COR_SUBTIT)
    secoes = [
        ("1. Resumen del Cambio",
         "Descripcion clara en 2-3 lineas: que cambia y por que."),
        ("2. Clasificacion",
         "Tipo (Bug Fix / Nueva Funcionalidad / Refactoring...), urgencia y complejidad."),
        ("3. Contexto y Motivacion",
         "AS-IS (situacion actual), problema identificado, TO-BE (situacion deseada)."),
        ("4. Detalle del Cambio",
         "Tabla AS-IS -> TO-BE para cada regla/comportamiento afectado. "
         "Tambien lista lo que NO cambia."),
        ("5. Impacto en Documentos",
         "Que documentos (PRD, Arquitectura, Spec, Plan, Deploy Guide) son afectados "
         "y que secciones hay que actualizar."),
        ("6. Impacto en el Codigo",
         "Lista de archivos a modificar/crear y descripcion de cada migration necesaria."),
        ("7. Tareas de Implementacion",
         "Tabla con ID (CR-T-01, CR-T-02...), descripcion, dependencias y criterio Done When."),
        ("8. Criterios de Aceptacion",
         "Checklist de comportamientos esperados despues del cambio."),
        ("9. Riesgos y Efectos Colaterales",
         "Riesgos identificados, probabilidad, impacto y plan de mitigacion."),
        ("10. Plan de Rollback",
         "Como revertir el codigo (git revert), la migration (alembic downgrade) "
         "y las variables de entorno, si es necesario."),
    ]

    t = doc.add_table(rows=len(secoes), cols=2)
    t.style = "Table Grid"
    for i, (secao, desc) in enumerate(secoes):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(secao); r0.bold = True; r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(desc); r1.font.size = Pt(10)

    add_spacer(doc)
    add_heading(doc, "Ejemplo de uso", 3, COR_SUBTIT)
    doc.add_paragraph(
        "El desarrollador nota que el listado de gastos necesita mostrar totales por estado "
        "(Pagado, Pendiente, Atrasado). Antes de tocar el codigo:"
    )
    passos_ex = [
        "Crea docs/changes/CR-004-totales-por-estado.md",
        "Completa AS-IS: el listado solo muestra el total general",
        "Completa TO-BE: el listado muestra total general + subtotales por estado",
        "Identifica que 03-SPEC.md, 04-IMPLEMENTATION-PLAN.md y el frontend necesitan actualizarse",
        "Lista las tareas: actualizar backend (servicio + endpoint), actualizar frontend (componente de tabla)",
        "Recien entonces comienza la implementacion, tarea por tarea",
    ]
    for i, passo in enumerate(passos_ex, 1):
        p = doc.add_paragraph(f"{i}. {passo}", style="List Number")
        p.runs[0].font.size = Pt(10)

    add_spacer(doc)


def add_template_prd(doc):
    add_heading(doc, "4.2  PRD - Product Requirements Document", 2, COR_PRD)

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_PRD)
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/01-template-prd.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, BRANCO)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_heading(doc, "Proposito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "El PRD define QUE construir: los requisitos del producto desde la perspectiva del usuario "
        "y del negocio. No describe como implementarlo; eso queda para la Arquitectura y la Spec. "
        "Es el documento que justifica la existencia de cada funcionalidad."
    )

    add_heading(doc, "Cuando crearlo", 3, COR_SUBTIT)
    items = [
        "Inicio de un nuevo proyecto",
        "Incorporacion de un modulo nuevo grande (ej: modulo de pagos, modulo de reportes)",
        "Cuando un CR identifica impacto en el PRD (nuevas user stories o requisitos)",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "Secciones del Template", 3, COR_SUBTIT)
    secoes = [
        ("1. Vision General",         "Descripcion del producto, problema que resuelve y publico objetivo"),
        ("2. Objetivos y Metricas",   "KPIs y metas medibles para cada objetivo"),
        ("3. Personas",               "Perfiles de los usuarios principales con necesidades y frustraciones"),
        ("4. Requisitos Funcionales", "Tabla RF-XXX con prioridad y detalle por bullet points"),
        ("5. Requisitos No Funcionales", "Performance, seguridad, accesibilidad (RNF-XXX)"),
        ("6. User Stories",           "US-XXX: Como [persona], quiero [accion], para que [beneficio]"),
        ("7. Reglas de Negocio",      "RN-XXX: restricciones y reglas que el sistema debe respetar"),
        ("8. Fuera de Alcance",       "Listar explicitamente lo que NO se construira en esta fase"),
        ("9. Dependencias",           "Sistemas externos, APIs de terceros, supuestos asumidos"),
        ("10. Glosario",              "Definicion de los terminos del dominio de negocio"),
        ("Apendice: Roadmap",         "Funcionalidades planificadas para fases futuras"),
    ]

    t = doc.add_table(rows=len(secoes), cols=2)
    t.style = "Table Grid"
    for i, (s, d) in enumerate(secoes):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(s); r0.bold = True; r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(d); r1.font.size = Pt(10)

    add_spacer(doc)
    add_info_box(
        doc,
        "Consejo: Mantene los IDs de los requisitos (RF-001, RN-001, US-001) consistentes en "
        "todos los documentos. La Spec va a referenciar RF-XXX; el Plan va a referenciar RF-XXX "
        "en las tareas. Eso genera trazabilidad completa de punta a punta.",
        COR_DESTAQUE
    )
    add_spacer(doc)


def add_template_arquitectura(doc):
    add_heading(doc, "4.3  Arquitectura", 2, COR_ARQ)

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_ARQ)
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/02-template-architecture.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, BRANCO)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_heading(doc, "Proposito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "La Arquitectura define COMO se construye el sistema: stack tecnologico, estructura de "
        "carpetas, modelado de datos, patrones y convenciones, estrategia de deploy y decisiones "
        "tecnicas. Es el documento de referencia para cualquier duda sobre 'como hacemos esto'."
    )

    add_heading(doc, "Cuando crearlo / actualizarlo", 3, COR_SUBTIT)
    items = [
        "Inicio del proyecto (definiendo todo el stack)",
        "Al incorporar una nueva tecnologia o libreria importante",
        "Cuando un CR introduce una nueva entidad de base de datos (actualizar diagrama ER)",
        "Al establecer un nuevo patron o convencion que debe seguir el equipo",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "Secciones del Template", 3, COR_SUBTIT)
    secoes = [
        ("1. Stack Tecnologico",       "Tabla: capa / tecnologia / version / justificacion"),
        ("2. Arquitectura General",    "Diagrama Mermaid del flujo entre componentes"),
        ("3. Estructura de Carpetas",  "Arbol completo de directorios del proyecto"),
        ("4. Modelado de Datos",       "Diagrama ER en Mermaid + detalle de cada entidad"),
        ("5. Patrones y Convenciones", "Nomenclatura, Git (branches/commits), reglas de API"),
        ("6. Integraciones Externas",  "Servicios de terceros con autenticacion y links a docs"),
        ("7. Estrategia de Tests",     "Tipos de test, herramientas y cobertura minima"),
        ("8. ADRs",                    "Registro de decisiones tecnicas con contexto y alternativas descartadas"),
        ("9. Deploy e Infraestructura","Pipeline CI/CD, health check, rollback"),
        ("10. Gestion de Dependencias","Politica de pinning de versiones y proceso de auditoria"),
    ]

    t = doc.add_table(rows=len(secoes), cols=2)
    t.style = "Table Grid"
    for i, (s, d) in enumerate(secoes):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(s); r0.bold = True; r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(d); r1.font.size = Pt(10)

    add_spacer(doc)
    add_info_box(
        doc,
        "ADRs (Architecture Decision Records): Cada vez que tomes una decision tecnica no obvia "
        "(ej: 'Por que usamos SQLite en dev en vez de PostgreSQL?'), registrala como ADR-XXX "
        "en la seccion 8. Incluye las alternativas que se descartaron y por que. "
        "Asi evitas que el equipo reabra discusiones ya resueltas.",
        COR_DESTAQUE
    )
    add_spacer(doc)


def add_template_spec(doc):
    add_heading(doc, "4.4  Especificacion Tecnica", 2, COR_SPEC)

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_SPEC)
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/03-template-spec.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, BRANCO)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_heading(doc, "Proposito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "La Spec es el documento mas tecnico del flujo. Detalla EXACTAMENTE lo que el desarrollador "
        "(o agente de IA) debe implementar: contratos de API, tipos/interfaces, logica de negocio "
        "paso a paso, schema de base de datos, validaciones, casos limite y plan de tests. "
        "No debe haber ambiguedad: si algo no esta en la Spec, no esta en el alcance."
    )

    add_heading(doc, "Cuando crearlo / actualizarlo", 3, COR_SUBTIT)
    items = [
        "Antes de implementar cualquier nueva funcionalidad",
        "Cuando un CR cambia contratos de API o logica de negocio existente",
        "Para documentar retroactivamente funcionalidades implementadas sin spec",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "Secciones del Template", 3, COR_SUBTIT)
    secoes = [
        ("1. Resumen de Cambios",          "Vision general y alcance de la iteracion"),
        ("Contratos de API (Vision General)", "Tabla macro de todos los endpoints: metodo, ruta, auth, body, respuesta"),
        ("2. Detalle Tecnico por Feature",  "Para cada RF-XXX: archivos, interfaces/types, logica de negocio, endpoints, base de datos, validaciones"),
        ("3. Componentes de UI",            "Props, estados (loading, empty, error) y comportamientos de los componentes"),
        ("4. Flujos Criticos",              "Diagramas secuenciales Mermaid de los flujos principales"),
        ("5. Casos Limite",                 "Tabla de escenarios no obvios y comportamiento esperado"),
        ("6. Plan de Tests",               "IDs de test (BT-XXX para backend, FT-XXX para flujo completo)"),
        ("7. Checklist de Implementacion", "Lista de verificacion para confirmar completitud"),
    ]

    t = doc.add_table(rows=len(secoes), cols=2)
    t.style = "Table Grid"
    for i, (s, d) in enumerate(secoes):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(s); r0.bold = True; r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(d); r1.font.size = Pt(10)

    add_spacer(doc)
    add_info_box(
        doc,
        "Consejo - Tabla de Contratos de API: Antes de detallar cada feature, incluye una tabla "
        "resumen con todos los endpoints del modulo (metodo, ruta, autenticacion, cuerpo, respuesta). "
        "Da una vision macro al desarrollador antes de entrar en los detalles.\n\n"
        "Patron PATCH: Usa siempre PATCH (no PUT) para actualizaciones parciales. "
        "En el backend, aplica exclude_unset=True (Pydantic) para ignorar campos no enviados.",
        COR_DESTAQUE
    )
    add_spacer(doc)


def add_template_plan(doc):
    add_heading(doc, "4.5  Plan de Implementacion", 2, COR_PLANO)

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COR_PLANO)
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/04-template-implementation-plan.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, BRANCO)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_heading(doc, "Proposito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "El Plan de Implementacion traduce los requisitos de la Spec en tareas concretas y "
        "secuenciadas. Define el orden de ejecucion, las dependencias entre tareas y el criterio "
        "de cierre (Done When) para cada una. Es el roadmap tecnico del desarrollador."
    )

    add_heading(doc, "Estructura de Tareas", 3, COR_SUBTIT)
    doc.add_paragraph(
        "Las tareas se organizan en grupos (1-6) y cada tarea tiene un ID unico (T-XXX). "
        "Los CRs agregados despues del plan inicial reciben secciones propias con IDs CR-T-XX."
    )

    grupos = [
        ("Grupo 1", "Setup e Infraestructura", "T-001 a T-005 - configuracion inicial del proyecto"),
        ("Grupo 2", "Modelos y Entidades",      "T-006 a T-008 - migrations, models, types"),
        ("Grupo 3", "Logica de Negocio",        "T-009 a T-013 - services y validaciones"),
        ("Grupo 4", "API / Controllers",         "T-014 a T-021 - endpoints y middlewares"),
        ("Grupo 5", "UI / Frontend",             "T-022 a T-028 - componentes y paginas"),
        ("Grupo 6", "Tests y Refinamiento",      "T-029 a T-033 - tests de integracion y E2E"),
    ]

    t = doc.add_table(rows=len(grupos), cols=3)
    t.style = "Table Grid"
    for i, (grupo, nome, desc) in enumerate(grupos):
        c0, c1, c2 = t.cell(i, 0), t.cell(i, 1), t.cell(i, 2)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, COR_PLANO); set_cell_bg(c1, bg); set_cell_bg(c2, bg)
        r0 = c0.paragraphs[0].add_run(grupo); r0.bold = True; r0.font.size = Pt(10); set_run_color(r0, BRANCO)
        r1 = c1.paragraphs[0].add_run(nome); r1.bold = True; r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(desc); r2.font.size = Pt(10)

    add_spacer(doc)
    add_heading(doc, "Columna Done When", 3, COR_SUBTIT)
    doc.add_paragraph(
        "Cada tarea tiene un criterio objetivo de cierre en la columna 'Done When'. "
        "Ejemplos validos: 'Migration corre sin error', 'Endpoint devuelve 201 con body correcto', "
        "'Componente se renderiza sin errores en consola'. Esto elimina la subjetividad sobre "
        "lo que significa 'terminar' una tarea."
    )

    add_spacer(doc)
    add_info_box(
        doc,
        "CRs en el Plan: Cuando un CR es aprobado, agrega una nueva seccion ## CR-XXX al final "
        "del archivo 04-IMPLEMENTATION-PLAN.md. Usa la misma estructura de grupos y tareas, "
        "pero con IDs CR-T-01, CR-T-02, etc. Asi se mantiene todo el historial de implementacion "
        "en un unico documento.",
        COR_DESTAQUE
    )
    add_spacer(doc)


def add_template_claude(doc):
    add_heading(doc, "4.6  CLAUDE.md - Instrucciones para el Agente de IA", 2, "5D3FD3")

    tbl = doc.add_table(rows=1, cols=1)
    set_table_no_spacing(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "5D3FD3")
    p = cell.paragraphs[0]
    r = p.add_run("Template: docs/templates/CLAUDE-template.md")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    set_run_color(r, BRANCO)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_spacer(doc)
    add_heading(doc, "Proposito", 3, COR_SUBTIT)
    doc.add_paragraph(
        "El CLAUDE.md es el archivo de instrucciones persistentes para agentes de IA (como Claude). "
        "Define la identidad del proyecto, el flujo de trabajo que debe seguirse, las reglas de "
        "implementacion, el stack tecnologico, el contexto actual y el historial de problemas "
        "conocidos. Cada nueva conversacion con el agente empieza con este archivo cargado "
        "automaticamente."
    )

    add_heading(doc, "Que incluir", 3, COR_SUBTIT)
    items = [
        ("Identidad del Proyecto",        "Nombre, descripcion resumida, stack y URL del repositorio"),
        ("Flujo de Desarrollo",           "Tabla de fases + Regla de Oro + instruccion de CR obligatorio"),
        ("Templates de Documentos",       "Tabla: tipo de doc -> ruta del template"),
        ("Reglas de Implementacion",      "Antes de Codear (5 reglas), Done When Universal, formato de commits"),
        ("Reglas para Cambios",           "Flujo CR en 8 pasos"),
        ("Estructura de Carpetas",        "Arbol del proyecto (sincronizado con el codigo real)"),
        ("Convenciones de Codigo",        "Nomenclatura por tipo de archivo y funcion"),
        ("Stack Tecnologico",             "Tabla completa con versiones"),
        ("Contexto Actual",               "CRs activos, ultima tarea implementada, documentos existentes"),
        ("Recordatorios Importantes",     "Reglas de comportamiento del agente (no inventar, un paso por vez, etc.)"),
        ("Troubleshooting",              "Problemas ya resueltos: dependencias, base de datos, entorno"),
    ]

    t = doc.add_table(rows=len(items), cols=2)
    t.style = "Table Grid"
    for i, (s, d) in enumerate(items):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(s); r0.bold = True; r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(d); r1.font.size = Pt(10)

    add_spacer(doc)
    add_info_box(
        doc,
        "Mantene el CLAUDE.md siempre actualizado. Al cerrar cada CR:\n"
        "1. Actualiza la seccion 'Contexto Actual' (CRs activos, ultima tarea)\n"
        "2. Actualiza la estructura de carpetas si se crearon nuevos archivos\n"
        "3. Agrega al Troubleshooting cualquier problema nuevo que se haya resuelto\n\n"
        "Un CLAUDE.md desactualizado hace que el agente tome decisiones basadas en contexto incorrecto.",
        COR_DESTAQUE
    )
    add_spacer(doc)


def add_guia_templates(doc):
    add_heading(doc, "4. Guia de Cada Template", 1)
    doc.add_paragraph(
        "Esta seccion presenta el proposito, cuando usar y las secciones de cada template "
        "disponible en docs/templates/. Siempre usa el template correspondiente al crear "
        "cualquier documento del flujo."
    )
    add_spacer(doc)
    add_template_cr(doc)
    add_template_prd(doc)
    add_template_arquitectura(doc)
    add_template_spec(doc)
    add_template_plan(doc)
    add_template_claude(doc)


# ─────────────────────────────────────────────
# 6. SISTEMA DE IDENTIFICADORES
# ─────────────────────────────────────────────
def add_identificadores(doc):
    add_heading(doc, "5. Sistema de Identificadores", 1)
    doc.add_paragraph(
        "El flujo SDD usa un sistema de IDs para crear trazabilidad entre documentos. "
        "Un ID en un documento debe aparecer exactamente igual en todos los otros documentos "
        "que hacen referencia a el."
    )

    ids = [
        ("RF-XXX",  COR_PRD,     "Requisito Funcional",             "PRD -> Spec -> Plan",           "RF-013: Modulo de Gastos Diarios"),
        ("RN-XXX",  COR_ARQ,     "Regla de Negocio",                "PRD -> Spec",                   "RN-019: Los gastos diarios no participan en la transicion de mes"),
        ("US-XXX",  COR_SPEC,    "User Story",                      "PRD",                           "US-020: Como usuario, quiero registrar un gasto diario"),
        ("T-XXX",   COR_PLANO,   "Tarea de Implementacion",         "Plan de Implementacion",        "T-045: Crear modelo DailyExpense"),
        ("CR-XXX",  COR_CR,      "Change Request",                  "Todas las fases",               "CR-005: Gastos Diarios"),
        ("CR-T-XX", COR_CR,      "Tarea de un CR especifico",       "Plan (seccion del CR)",         "CR-T-01: Crear migration daily_expenses"),
        ("ADR-XXX", COR_ARQ,     "Architecture Decision Record",    "Arquitectura",                  "ADR-014: Uso de Alembic para migrations"),
        ("BT-XXX",  COR_VALIDAR, "Backend Test",                    "Spec -> Tests",                 "BT-011: POST /api/daily-expenses devuelve 201"),
        ("FT-XXX",  COR_VALIDAR, "Full Flow Test",                  "Spec -> Tests",                 "FT-030: Usuario registra gasto y lo ve en el listado"),
    ]

    t = doc.add_table(rows=1 + len(ids), cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(["Prefijo", "Tipo", "Usado en", "Ejemplo"]):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10); set_run_color(r, BRANCO)

    for i, (prefijo, cor, tipo, usado_en, ejemplo) in enumerate(ids):
        row_data = [prefijo, tipo, usado_en, ejemplo]
        for j, val in enumerate(row_data):
            cell = t.cell(i + 1, j)
            if j == 0:
                set_cell_bg(cell, cor)
                r = cell.paragraphs[0].add_run(val)
                r.bold = True; r.font.size = Pt(10); set_run_color(r, BRANCO)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
                set_cell_bg(cell, bg)
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(9 if j == 3 else 10)

    add_spacer(doc)
    add_info_box(
        doc,
        "Trazabilidad: Si el RF-013 aparece en el PRD, debe aparecer como referencia "
        "en la Spec (seccion de la feature) y en el Plan (columna Ref de las tareas). "
        "Eso permite rastrear cualquier requisito de punta a punta.",
        COR_DESTAQUE
    )
    add_spacer(doc)


# ─────────────────────────────────────────────
# 7. PROMPT DE INICIO DE PROYECTO
# ─────────────────────────────────────────────
def add_prompt_inicio(doc):
    add_heading(doc, "6. Prompt para Iniciar un Nuevo Proyecto", 1)

    doc.add_paragraph(
        "Usa el siguiente prompt al iniciar una nueva conversacion con un agente de IA (como Claude) "
        "para configurar el flujo SDD en un nuevo proyecto. Adapta los campos entre corchetes."
    )

    add_info_box(
        doc,
        "Atajo recomendado - Skill /feature (Claude Code)\n\n"
        "Si usas Claude Code con la skill /feature configurada, no tenes que copiar "
        "los prompts de abajo manualmente. Basta con invocar /feature en una nueva conversacion: "
        "la skill lee el CLAUDE.md, determina automaticamente el modo (Nuevo Proyecto o Change Request), "
        "crea los documentos, abre la branch, implementa, hace revision de seguridad, "
        "valida el build y realiza el commit/merge, todo sin intervencion manual.\n\n"
        "Los prompts a continuacion son utiles para entender el proceso en detalle o para ambientes "
        "donde la skill no esta configurada.",
        "E8DAEF"
    )
    add_spacer(doc)
    add_heading(doc, "Prompt de Setup Inicial", 3, COR_SUBTIT)

    prompt_texto = """\
Vamos a iniciar un nuevo proyecto de software llamado [NOMBRE DEL PROYECTO].

## Contexto del Proyecto
- **Descripcion:** [Lo que hace el sistema y para quien]
- **Stack:** [Frontend: ej. React + TypeScript] / [Backend: ej. FastAPI + Python] / [Base de datos: ej. PostgreSQL]
- **Repositorio:** [URL de Git o "crear nuevo"]

## Lo que necesito que hagas

**Paso 1 - Crear el CLAUDE.md:**
Leer el template en docs/templates/CLAUDE-template.md y crear el archivo CLAUDE.md
en la raiz del proyecto. Reemplaza todos los [placeholders] con la informacion de arriba.

**Paso 2 - Crear el PRD (01-PRD.md):**
Leer el template en docs/templates/01-template-prd.md y crear docs/01-PRD.md
con base en la siguiente informacion:
[DESCRIBIR LAS FUNCIONALIDADES PRINCIPALES DEL PRODUCTO]

**Paso 3 - Crear la Arquitectura (02-ARCHITECTURE.md):**
Con base en el PRD, crear docs/02-ARCHITECTURE.md usando el template
docs/templates/02-template-architecture.md.

**Paso 4 - Crear la Spec Tecnica (03-SPEC.md):**
Con base en el PRD y la Arquitectura, crear docs/03-SPEC.md usando el template
docs/templates/03-template-spec.md. Detallar los endpoints, tipos, logica
de negocio y plan de tests para cada requisito funcional.

**Paso 5 - Crear el Plan de Implementacion (04-IMPLEMENTATION-PLAN.md):**
Con base en la Spec, crear docs/04-IMPLEMENTATION-PLAN.md usando el template
docs/templates/04-template-implementation-plan.md.

## Reglas que tenes que seguir
- Documentar PRIMERO, codigo DESPUES
- Nunca inventar funcionalidades que no esten en el PRD
- Usar los IDs RF-XXX, RN-XXX, T-XXX, CR-XXX conforme al flujo SDD
- Al terminar, confirmar que los 5 documentos fueron creados\
"""

    add_code_box(doc, prompt_texto)
    add_spacer(doc)

    add_heading(doc, "Prompt para un Nuevo CR (Feature o Correccion)", 3, COR_SUBTIT)

    prompt_cr = """\
Necesito implementar un nuevo cambio en el proyecto [NOMBRE DEL PROYECTO].

## Descripcion del Cambio
[DESCRIBIR LO QUE HAY QUE HACER]

## Instrucciones
1. Leer el archivo CLAUDE.md para entender el contexto actual del proyecto
2. Crear un Change Request usando el template docs/templates/00-template-change-request.md
   - Guardar en docs/changes/CR-[SIGUIENTE NUMERO]-[slug].md
   - Completar todas las secciones: AS-IS, TO-BE, impacto en docs, tareas, criterios de aceptacion
3. Evaluar el impacto y actualizar los documentos afectados (PRD, Arquitectura, Spec, Plan)
4. Implementar las tareas del CR de a una por vez, confirmando cada Done When
5. Al terminar: verificar el build, escribir los tests, actualizar el CLAUDE.md, hacer el commit\
"""

    add_code_box(doc, prompt_cr)
    add_spacer(doc)

    add_info_box(
        doc,
        "Consejo: Se especifico en la descripcion del cambio. En vez de 'mejorar el listado de "
        "gastos', escribi 'agregar totales por estado (Pagado, Pendiente, Atrasado) en la tabla "
        "de gastos mensuales, debajo del total general'. Cuanto mas especifico, menos ambiguedad "
        "para que el agente resuelva.",
        COR_DESTAQUE
    )
    add_spacer(doc)
    add_info_box(
        doc,
        "Alternativa con skill: En Claude Code, usa /feature en lugar de este prompt. "
        "La skill ejecuta el flujo de CR completo de forma automatizada: creacion del documento CR, "
        "branch, implementacion, revision de seguridad, build y commit.",
        "E8DAEF"
    )
    add_spacer(doc)


# ─────────────────────────────────────────────
# 8. DONE WHEN UNIVERSAL
# ─────────────────────────────────────────────
def add_done_when(doc):
    add_heading(doc, "7. Checklist 'Done When' Universal", 1)

    doc.add_paragraph(
        "Toda tarea (T-XXX o CR-T-XX) se considera cerrada solo cuando todos los criterios "
        "obligatorios abajo se satisfacen. Usa esta checklist como revision antes de hacer el commit."
    )

    add_spacer(doc)
    add_heading(doc, "Obligatorios (todas las tareas)", 3, COR_SUBTIT)
    obrigatorios = [
        "Funcionalidad implementada conforme a lo descrito en la tarea",
        "La app corre localmente sin errores (backend + frontend)",
        "Los tests existentes siguen pasando (sin regresion)",
        "Nuevos tests cubren la funcionalidad agregada o modificada",
        "El commit sigue Conventional Commits y referencia el ID de la tarea",
    ]
    t_obrig = doc.add_table(rows=len(obrigatorios), cols=2)
    for i, item in enumerate(obrigatorios):
        c0, c1 = t_obrig.cell(i, 0), t_obrig.cell(i, 1)
        set_cell_bg(c0, "D5F5E3")
        r0 = c0.paragraphs[0].add_run("OK")
        r0.bold = True; r0.font.size = Pt(11)
        set_run_color(r0, "1E8449")
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c1, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r1 = c1.paragraphs[0].add_run(item)
        r1.font.size = Pt(11)

    add_spacer(doc)
    add_heading(doc, "Si aplica", 3, COR_SUBTIT)
    aplicaveis = [
        ("Migration",     "alembic upgrade head + alembic downgrade -1 testeados sin errores"),
        ("Endpoints",     "Responden con status codes correctos (201, 204, 400, 401, 404, etc.)"),
        ("Documentacion", "Documentos afectados actualizados (Spec, Arquitectura, CLAUDE.md)"),
        ("Frontend",      "Sin errores ni warnings en la consola del navegador"),
        ("Build",         "Build TypeScript pasa sin errores de tipo (npx tsc --noEmit)"),
    ]

    t_aplic = doc.add_table(rows=len(aplicaveis), cols=3)
    t_aplic.style = "Table Grid"
    for i, (contexto, item) in enumerate(aplicaveis):
        c0, c1, c2 = t_aplic.cell(i, 0), t_aplic.cell(i, 1), t_aplic.cell(i, 2)
        set_cell_bg(c0, "FEF9E7")
        set_cell_bg(c1, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        set_cell_bg(c2, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r0 = c0.paragraphs[0].add_run("Si aplica"); r0.font.size = Pt(9); r0.italic = True
        r1 = c1.paragraphs[0].add_run(contexto); r1.bold = True; r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(item); r2.font.size = Pt(10)

    add_spacer(doc)


# ─────────────────────────────────────────────
# 9. REFERENCIA RAPIDA
# ─────────────────────────────────────────────
def add_referencia_rapida(doc):
    add_heading(doc, "8. Referencia Rapida", 1)

    add_heading(doc, "Cuando crear cada documento", 3, COR_SUBTIT)

    situacoes = [
        ("Voy a cambiar algo que ya existe",       "Change Request (CR)",                 "00-template-change-request.md", COR_CR),
        ("Voy a crear un modulo nuevo",            "PRD + Arquitectura + Spec + Plan",    "01 a 04 templates",             COR_PRD),
        ("Voy a cambiar el stack o la estructura", "Arquitectura + CR",                   "02-template-architecture.md",   COR_ARQ),
        ("Voy a implementar una nueva feature",    "Spec + Plan (o CR si ya existe)",     "03-template-spec.md",           COR_SPEC),
        ("Voy a iniciar un nuevo proyecto",        "CLAUDE.md + todos los documentos",    "CLAUDE-template.md + 01-04",    "5D3FD3"),
        ("Voy a hacer deploy",                     "Deploy Guide",                        "05-DEPLOY-GUIDE.md",            "1A5276"),
        ("Quiero automatizar el pipeline SDD",     "Skill /feature (Claude Code)",        ".claude/skills/feature/",       "6C3483"),
    ]

    t = doc.add_table(rows=1 + len(situacoes), cols=3)
    t.style = "Table Grid"
    for j, h in enumerate(["Situacion", "Documento a Crear", "Template"]):
        cell = t.cell(0, j)
        set_cell_bg(cell, COR_TITULO)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10); set_run_color(r, BRANCO)

    for i, (sit, doc_crear, template, cor) in enumerate(situacoes):
        for j, val in enumerate([sit, doc_crear, template]):
            cell = t.cell(i + 1, j)
            bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
            if j == 1:
                set_cell_bg(cell, cor)
                r = cell.paragraphs[0].add_run(val)
                r.bold = True; r.font.size = Pt(10); set_run_color(r, BRANCO)
            else:
                set_cell_bg(cell, bg)
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(10)

    add_spacer(doc)

    add_heading(doc, "Conventional Commits", 3, COR_SUBTIT)

    commits = [
        ("feat:",     "Nueva feature implementada",                    "feat: implement CR-005 - daily expenses module"),
        ("fix:",      "Correccion de bug",                            "fix: CR-007 - evitar duplicacion de cuotas en transicion de mes"),
        ("docs:",     "Actualizacion de documentacion",               "docs: update spec and PRD for CR-005"),
        ("refactor:", "Refactorizacion sin cambio de comportamiento", "refactor: extract status logic to dedicated service"),
        ("test:",     "Agregado o correccion de tests",               "test: add BT-011 to BT-015 for daily expenses endpoints"),
        ("chore:",    "Tareas de mantenimiento",                      "chore: update dependencies, pin bcrypt==4.0.*"),
    ]

    t2 = doc.add_table(rows=len(commits), cols=3)
    t2.style = "Table Grid"
    for i, (prefijo, desc, ejemplo) in enumerate(commits):
        c0, c1, c2 = t2.cell(i, 0), t2.cell(i, 1), t2.cell(i, 2)
        set_cell_bg(c0, COR_TITULO)
        bg = COR_CINZA_BG if i % 2 == 0 else BRANCO
        set_cell_bg(c1, bg); set_cell_bg(c2, bg)
        r0 = c0.paragraphs[0].add_run(prefijo); r0.font.name = "Courier New"; r0.font.size = Pt(10); set_run_color(r0, "85C1E9")
        r1 = c1.paragraphs[0].add_run(desc); r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(ejemplo); r2.font.name = "Courier New"; r2.font.size = Pt(9)

    add_spacer(doc)

    add_heading(doc, "Flujo CR en 8 Pasos", 3, COR_SUBTIT)

    passos_cr = [
        "Crear el CR + branch: git checkout -b feat/CR-[XXX]-[slug]",
        "Completar todas las secciones: resumen, AS-IS/TO-BE, impacto, tareas, criterios de aceptacion",
        "Evaluar el impacto en los documentos: que secciones de PRD, Spec, Arquitectura, Plan son afectadas",
        "Actualizar los documentos afectados ANTES de implementar",
        "Implementar las tareas del CR de a una por vez",
        "Validar los criterios de aceptacion del CR + checklist Done When Universal",
        "Verificar el build: npx tsc --noEmit (frontend) + tests pasando (backend)",
        "Merge y push: git merge feat/CR-XXX --no-ff + git push origin master",
    ]

    t3 = doc.add_table(rows=len(passos_cr), cols=2)
    cores_passo = [COR_CR, COR_CR, COR_PRD, COR_ARQ, COR_SPEC, COR_PLANO, COR_VALIDAR, "1A5276"]
    for i, (passo, cor) in enumerate(zip(passos_cr, cores_passo)):
        c0, c1 = t3.cell(i, 0), t3.cell(i, 1)
        set_cell_bg(c0, cor)
        set_cell_bg(c1, COR_CINZA_BG if i % 2 == 0 else BRANCO)
        r0 = c0.paragraphs[0].add_run(str(i + 1))
        r0.bold = True; r0.font.size = Pt(14); set_run_color(r0, BRANCO)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = c1.paragraphs[0].add_run(passo); r1.font.size = Pt(11)

    add_spacer(doc, 2)
    add_info_box(
        doc,
        "Este manual esta disponible en: docs/manual-workflow-sdd-es.docx\n"
        "Los templates estan en: docs/templates/\n"
        "El template generico del CLAUDE.md esta en: docs/templates/CLAUDE-template.md",
        COR_CINZA_BG
    )
    add_spacer(doc)


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def main():
    doc = Document()
    configure_document(doc)

    add_portada(doc)
    add_introduccion(doc)
    add_vision_general(doc)
    add_page_break(doc)
    add_tipos_de_trabajo(doc)
    add_page_break(doc)
    add_guia_templates(doc)
    add_page_break(doc)
    add_identificadores(doc)
    add_page_break(doc)
    add_prompt_inicio(doc)
    add_page_break(doc)
    add_done_when(doc)
    add_referencia_rapida(doc)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, "manual-workflow-sdd-es.docx")
    doc.save(output_path)
    print(f"[OK]  Manual en espanol generado: {output_path}")


if __name__ == "__main__":
    main()
